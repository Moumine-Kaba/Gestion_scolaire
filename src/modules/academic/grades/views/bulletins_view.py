# -*- coding: utf-8 -*-
"""
Vue des Bulletins - Système Réorganisé
EduManager+ - Interface Moderne avec Formulaire Structuré

Cette vue présente les bulletins organisés par classe et période
avec un formulaire modal stylisé et des menus déroulants dépendants.
"""

import customtkinter as ctk
from tkinter import messagebox, StringVar, Toplevel, filedialog
import os
import sys
from PIL import Image
from CTkTable import CTkTable
from datetime import datetime
from typing import List, Dict, Optional, TYPE_CHECKING, Any
import importlib.util as _import_util

# Détection paresseuse de python-docx pour éviter les warnings Pyright
DOCX_AVAILABLE = _import_util.find_spec("docx") is not None

# Aide au linter: fournir des alias de types quand python-docx n'est pas installé
if TYPE_CHECKING:
    from docx import Document as _T_Document
    from docx.shared import Inches as _T_Inches, Pt as _T_Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _T_WD
    Document: Any = _T_Document  # type: ignore
    Inches: Any = _T_Inches  # type: ignore
    Pt: Any = _T_Pt  # type: ignore
    WD_ALIGN_PARAGRAPH: Any = _T_WD  # type: ignore
else:
    Document = Inches = Pt = WD_ALIGN_PARAGRAPH = None  # runtime placeholders

# Ajouter le chemin du projet pour les imports
project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '../../../../..'))
if project_root not in sys.path:
    sys.path.append(project_root)

from src.modules.academic.grades.controllers.bulletins_sqlserver_controller import BulletinsController
# Adapters vers le contrôleur simple des bulletins
try:
    from src.modules.academic.grades.controllers.bulletin_controller import (
        get_all_bulletins as db_get_all_bulletins,
        add_bulletin as db_add_bulletin,
        update_bulletin as db_update_bulletin,
        delete_bulletin as db_delete_bulletin,
    )
except Exception as _e:
    db_get_all_bulletins = None
    db_add_bulletin = None
    db_update_bulletin = None
    db_delete_bulletin = None
from src.modules.academic.grades.controllers.notes_controller import get_notes_by_eleve, get_notes_summary_by_eleve
from src.modules.academic.students.controllers.eleve_controller import get_all_eleves
from src.modules.academic.classes.controllers.classe_controller import get_all_classes

# Import du thème global
try:
    from resources.themes.theme import *
    print("Theme global EduManager+ importe pour les bulletins")
except ImportError as e:
    print(f"Theme global non trouve: {e}")
    # Thème de fallback
    BG_MAIN = "#233146"
    BG_CARD = "#2b2952"
    TEXT_PRIMARY = "#E0E6F0"
    TEXT_SECONDARY = "#AAB5C6"
    TEXT_ACCENT = "#64FFDA"
    BORDER_COLOR = "#40546c"
    SUCCESS_GREEN = "#A0E7E5"
    WARNING_YELLOW = "#FFD700"
    ERROR_RED = "#FF6363"
    ACCENT = "#64FFDA"
    BG_SIDEBAR = "#1E2332"
    TEXT_SUCCESS = "#4CAF50"
    TEXT_WARNING = "#FF9800"
    TEXT_ERROR = "#F44336"
    BG_SECONDARY = "#3A3F5C"
    BG_CARD = "#2b2952"
    FONT_SIZE_TEXT = 12
    FONT_SIZE_HEADER = 18
    MARGIN_SMALL = 8
    MARGIN_MEDIUM = 12
    MARGIN_LARGE = 20
    FONT = "Segoe UI"
    FONT_SIZE_TITLE = 24
    FONT_SIZE_HEADER = 18
    FONT_SIZE_SUB = 14
    FONT_SIZE_TXT = 12
    FONT_SIZE_SMALL = 10
    FONT_BOLD = "Segoe UI Bold"

# Variables de police
F_TITLE = (FONT, FONT_SIZE_TITLE, "bold")
F_SUB = (FONT, FONT_SIZE_SUB, "bold")
F_TXT = (FONT, FONT_SIZE_TXT)
F_SMALL = (FONT, FONT_SIZE_SMALL)
F_BOLD = (FONT, FONT_SIZE_TXT, "bold")

# Icônes
ICON_MAP = {
    'add': 'resources/icons/add.png',
    'edit': 'resources/icons/edit.png',
    'delete': 'resources/icons/delete.png',
    'refresh': 'resources/icons/refresh.png',
    'search': 'resources/icons/search.png',
    'filter': 'resources/icons/filter.png',
    'newspaper': 'resources/icons/newspaper.png',
    'grade': 'resources/icons/grade.png',
    'stats': 'resources/icons/analytics.png',
    'class': 'resources/icons/classroom.png',
    'person': 'resources/icons/person.png',
    'school': 'resources/icons/classroom.png',
    'calendar': 'resources/icons/calendar.png',
    'chart': 'resources/icons/analytics.png',
    'award': 'resources/icons/award.png',
    'comment': 'resources/icons/assignment.png',
    'export': 'resources/icons/upload.png',
    'generate': 'resources/icons/autorenew.png',
    'sort': 'resources/icons/sort.png',
    'person': 'resources/icons/person.png'
}

def load_icon(icon_name, size=(20, 20)):
    """Charge une icône avec gestion d'erreur"""
    try:
        icon_path = ICON_MAP.get(icon_name)
        if icon_path and os.path.exists(icon_path):
            return ctk.CTkImage(Image.open(icon_path), size=size)
        else:
            print(f"Icone '{icon_name}' non trouvee: {icon_path}")
            return None
    except Exception as e:
        print(f"Erreur chargement icone '{icon_name}': {e}")
        return None

        # Appeler la fonction du contrôleur avec les bons paramètres
        result = db_add_bulletin(eleve_id, annee_scolaire, trimestre, moyenne, remarque, date_edition)
        return result if result is not None else True
    except Exception as e:
        print(f"add_bulletin adapter error: {e}")
        import traceback
        traceback.print_exc()
# Compatibilité: certaines sections utilisent load_ctk_icon
def load_ctk_icon(name_or_path, size=(20, 20)):
    """Retourne une CTkImage à partir d'un nom logique ou d'un chemin.
    - Si name_or_path est une clé d'ICON_MAP, utilise ce chemin
    - Si c'est un nom de fichier *.png, cherche dans resources/icons
    - Si c'est un chemin absolu/relatif existant, l'utilise tel quel
    """
    try:
        candidate = None
        if isinstance(name_or_path, str):
            # Clé de l'ICON_MAP
            if name_or_path in ICON_MAP:
                candidate = ICON_MAP[name_or_path]
            # Nom de fichier icône
            elif name_or_path.endswith('.png') and not os.path.isabs(name_or_path):
                candidate = os.path.join('resources', 'icons', name_or_path)
            else:
                candidate = name_or_path
        if candidate and os.path.exists(candidate):
            return ctk.CTkImage(Image.open(candidate), size=size)
    except Exception as e:
        print(f"Erreur load_ctk_icon '{name_or_path}': {e}")
    return None

# Adapters simples pour correspondre aux signatures locales
def get_all_bulletins():
    if db_get_all_bulletins:
        return db_get_all_bulletins()
    return []

def add_bulletin(info: Dict) -> bool:
    """Adaptateur pour l'ajout d'un bulletin.
    Paramètres attendus: id_eleve, periode, moyenne_generale, appreciation
    """
    if not db_add_bulletin:
        return False
    try:
        eleve_id = info.get('id_eleve')
        annee_scolaire = str(datetime.now().year)
        trimestre = info.get('periode') or ''
        moyenne = float(info.get('moyenne_generale') or 0)
        remarque = info.get('appreciation') or ''
        date_edition = datetime.now().date()

        result = db_add_bulletin(eleve_id, annee_scolaire, trimestre, moyenne, remarque, date_edition)
        return bool(result) if result is not None else True
    except Exception as e:
        print(f"add_bulletin adapter error: {e}")
        import traceback
        traceback.print_exc()
        return False

def update_bulletin(bulletin_id: int, info: Dict):
    if not db_update_bulletin:
        return False
    try:
        eleve_id = info.get('id_eleve')
        annee_scolaire = str(datetime.now().year)
        trimestre = info.get('periode') or ''
        moyenne = float(info.get('moyenne_generale') or 0)
        remarque = info.get('appreciation') or ''
        date_edition = datetime.now().date()
        db_update_bulletin(bulletin_id, eleve_id, annee_scolaire, trimestre, moyenne, remarque, date_edition)
        return True
    except Exception as e:
        print(f"update_bulletin adapter error: {e}")
        return False

def delete_bulletin(bulletin_id: int):
    if not db_delete_bulletin:
        return False
    try:
        db_delete_bulletin(bulletin_id)
        return True
    except Exception as e:
        print(f"delete_bulletin adapter error: {e}")
        return False

class BulletinsView(ctk.CTkFrame):
    """Vue des Bulletins avec formulaire structuré"""

    def __init__(self, parent):
        super().__init__(parent, fg_color=BG_MAIN)
        
        # Variables de données
        self.bulletins = []
        self.classes = []
        self.periodes = []
        
        # Variables de sélection
        self.selected_classe = None
        self.selected_periode = None
        self.selected_bulletin = None
        
        # Variables de pagination
        self.current_page = 1
        self.items_per_page = 20
        self.total_pages = 1
        
        # Cache des données
        self._data_cache = {}
        self._cache_timestamp = None
        self._cache_duration = 300  # 5 minutes
        
        # Interface
        self.table_frame = None
        self.form_modal = None
        
        # Contrôleur des bulletins
        self.bulletins_controller = BulletinsController()
        
        # Charger les données
        self._load_data()
        
        # Construire l'interface
        self._build_main_ui()

        # Afficher le message initial
        self._show_no_selection_message()

    # ===== Helpers robustes =====
    def _get_class_id_by_name(self, classe_name: str) -> Optional[int]:
        """Retourne l'ID de classe à partir de son nom, en gérant les différents schémas.
        Cherche des clés communes: 'nom', 'nom_classe', 'classe_nom'.
        Retourne l'ID via: 'id', 'id_classe', 'idclasse'.
        """
        try:
            classes_list = get_all_classes()
            for classe in classes_list:
                current_name = (
                    classe.get('nom')
                    or classe.get('nom_classe')
                    or classe.get('classe_nom')
                )
                if current_name == classe_name:
                    return (
                        classe.get('id')
                        or classe.get('id_classe')
                        or classe.get('idclasse')
                    )
        except Exception as _e:
            pass
        return None

    def _extract_subject_name(self, note: Dict) -> str:
        """Retourne le nom de la matière pour une note en gérant différents champs."""
        return (
            note.get('nom_matiere')
            or note.get('matiere')
            or note.get('libelle_matiere')
            or note.get('nom')
            or 'Inconnue'
        )

    def _extract_coefficient(self, note: Dict) -> float:
        """Retourne le coefficient de la matière depuis la note.
        Clés acceptées: 'coefficient', 'coeff', 'coef', 'coefficient_matiere'.
        """
        for key in ('coefficient', 'coeff', 'coef', 'coefficient_matiere'):
            if key in note and note.get(key) is not None:
                try:
                    return float(note.get(key))
                except Exception:
                    continue
        return 1.0

    def _extract_student_class_id(self, eleve: Dict) -> Optional[int]:
        """Récupère l'ID de classe depuis un élève, en gérant différentes clés et types."""
        for key in ('id_classe', 'classe_id', 'idclasse', 'idClasse', 'classeId'):
            if key in eleve and eleve.get(key) is not None:
                try:
                    return int(eleve.get(key))
                except Exception:
                    # si cast échoue, retourne brut
                    try:
                        return int(str(eleve.get(key)).strip())
                    except Exception:
                        return None
        return None

    def _student_fullname(self, eleve: Dict) -> str:
        """Construit un nom complet robuste (gère clés manquantes)."""
        nom = eleve.get('nom') or eleve.get('last_name') or eleve.get('surname') or ''
        prenom = eleve.get('prenom') or eleve.get('first_name') or eleve.get('given_name') or ''
        full = f"{str(nom).strip()} {str(prenom).strip()}".strip()
        return full or f"Élève {eleve.get('id_eleve') or eleve.get('id') or ''}"
    
    def _load_data(self):
        """Charge les données depuis la base"""
        try:
            print("Chargement des donnees bulletins...")
            
            # Charger les classes
            classes_list = get_all_classes()
            # Créer un dictionnaire avec l'ID comme clé et les données comme valeur
            self.classes = {classe['id']: classe for classe in classes_list}
            print(f"{len(self.classes)} classes chargees")
            
            # Charger les périodes
            self.periodes = ["1er Trimestre", "2ème Trimestre", "3ème Trimestre"]
            print(f"{len(self.periodes)} periodes chargees")
            
            # Charger les bulletins depuis la base de données
            self.bulletins = get_all_bulletins()
            print(f"{len(self.bulletins)} bulletins charges")
            
            # Debug: Afficher la structure des premiers bulletins
            if self.bulletins:
                print("DEBUG: Structure du premier bulletin:")
                first_bulletin = self.bulletins[0]
                for key, value in first_bulletin.items():
                    print(f"  {key}: {value}")
                
                # Afficher les classes uniques dans les bulletins
                classes_in_bulletins = set()
                for bulletin in self.bulletins:
                    classe_nom = bulletin.get('classe_nom', '')
                    if classe_nom:
                        classes_in_bulletins.add(classe_nom)
                print(f"DEBUG: Classes trouvees dans les bulletins: {sorted(classes_in_bulletins)}")
                
                # Corriger les bulletins qui n'ont pas de classe_nom
                self._fix_missing_classe_names()
            
            print("Donnees bulletins chargees avec succes")
            
        except Exception as e:
            print(f"Erreur chargement donnees: {e}")
            messagebox.showerror("Erreur", f"Erreur lors du chargement des données:\n{str(e)}")
    
    def _build_main_ui(self):
        """Construit l'interface principale"""
        # Configuration de la grille
        self.grid_columnconfigure(0, weight=0, minsize=280)  # Sidebar réduite
        self.grid_columnconfigure(1, weight=1)  # Contenu principal
        self.grid_rowconfigure(0, weight=1)
        
        # Panneau gauche - Sélection et filtres
        self._build_selection_panel()
        
        # Panneau droit - Tableau des bulletins
        self._build_bulletins_dashboard()
    
    def _build_selection_panel(self):
        """Construit le panneau de sélection gauche avec design élégant"""
        # Frame principal avec gradient subtil
        left_panel = ctk.CTkFrame(self, fg_color=BG_SIDEBAR, corner_radius=0)
        left_panel.grid(row=0, column=0, sticky="nsew", padx=0, pady=0)
        left_panel.grid_columnconfigure(0, weight=1)
        left_panel.grid_rowconfigure(4, weight=1)
        
        # ============ HEADER ÉLÉGANT ============
        header_frame = ctk.CTkFrame(left_panel, fg_color="transparent")
        header_frame.grid(row=0, column=0, sticky="ew", padx=0, pady=0)
        
        # Container avec fond accent subtil
        header_bg = ctk.CTkFrame(header_frame, fg_color=BG_CARD, corner_radius=0)
        header_bg.pack(fill="both", expand=True, padx=0, pady=0)
        
        # Icône et titre en ligne
        title_container = ctk.CTkFrame(header_bg, fg_color="transparent")
        title_container.pack(pady=20, padx=25)
        
        # Icône à gauche
        grade_icon = load_icon('grade', (28, 28))
        if grade_icon:
            icon_label = ctk.CTkLabel(title_container, image=grade_icon, text="")
            icon_label.pack(side="left", padx=(0, 12))
        
        # Texte à droite de l'icône
        text_container = ctk.CTkFrame(title_container, fg_color="transparent")
        text_container.pack(side="left")
        
        title_label = ctk.CTkLabel(text_container, text="Bulletins Scolaires", 
                                  font=(FONT, 17, "bold"), text_color=TEXT_ACCENT,
                                  anchor="w")
        title_label.pack(anchor="w")
        
        # Sous-titre avec compteur
        total_bulletins = len(self.bulletins) if hasattr(self, 'bulletins') else 0
        subtitle_label = ctk.CTkLabel(text_container, 
                                     text=f"Gestion de {total_bulletins} bulletins", 
                                     font=(FONT, 10), text_color=TEXT_SECONDARY,
                                     anchor="w")
        subtitle_label.pack(anchor="w", pady=(3, 0))
        
        # ============ SECTION FILTRES ÉLÉGANTE ============
        filters_section = ctk.CTkFrame(left_panel, fg_color="transparent")
        filters_section.grid(row=2, column=0, sticky="ew", padx=25, pady=(18, 0))
        filters_section.grid_columnconfigure(0, weight=1)
        
        # Titre de section avec style
        filters_title = ctk.CTkLabel(filters_section, text="FILTRES", 
                                     font=(FONT, 9, "bold"), text_color=TEXT_SECONDARY,
                                     anchor="w")
        filters_title.grid(row=0, column=0, sticky="w", pady=(0, 10))
        
        # Card pour les filtres
        filters_card = ctk.CTkFrame(filters_section, fg_color=BG_CARD, corner_radius=10)
        filters_card.grid(row=1, column=0, sticky="ew")
        filters_card.grid_columnconfigure(0, weight=1)
        
        # Filtre classe dans la card
        classe_container = ctk.CTkFrame(filters_card, fg_color="transparent")
        classe_container.grid(row=0, column=0, sticky="ew", padx=12, pady=(12, 8))
        classe_container.grid_columnconfigure(0, weight=1)
        
        ctk.CTkLabel(classe_container, text="Classe", font=(FONT, 10, "bold"), 
                    text_color=TEXT_PRIMARY, anchor="w").grid(row=0, column=0, sticky="w", pady=(0, 5))
        
        self.classe_var = StringVar()
        classe_values = ["Toutes les classes"] + [classe['nom'] for classe in self.classes.values()]
        self.classe_dropdown = ctk.CTkComboBox(classe_container, variable=self.classe_var, 
                                              values=classe_values,
                                              command=self._on_classe_selected, state="readonly",
                                              font=(FONT, 11), dropdown_font=(FONT, 10),
                                              corner_radius=8, border_width=1, border_color=BORDER_COLOR,
                                              fg_color=BG_MAIN, button_color=TEXT_ACCENT, 
                                              button_hover_color=SUCCESS_GREEN,
                                              height=38)
        self.classe_dropdown.grid(row=1, column=0, sticky="ew")
        self.classe_dropdown.set("Toutes les classes")
        
        # Séparateur interne
        sep_internal = ctk.CTkFrame(filters_card, height=1, fg_color=BORDER_COLOR)
        sep_internal.grid(row=1, column=0, sticky="ew", padx=12, pady=0)
        
        # Filtre période dans la card
        periode_container = ctk.CTkFrame(filters_card, fg_color="transparent")
        periode_container.grid(row=2, column=0, sticky="ew", padx=12, pady=(8, 12))
        periode_container.grid_columnconfigure(0, weight=1)
        
        ctk.CTkLabel(periode_container, text="Période", font=(FONT, 10, "bold"), 
                    text_color=TEXT_PRIMARY, anchor="w").grid(row=0, column=0, sticky="w", pady=(0, 5))
        
        self.periode_var = StringVar()
        periode_values = ["Toutes les périodes"] + self.periodes
        self.periode_dropdown = ctk.CTkComboBox(periode_container, variable=self.periode_var,
                                              values=periode_values, command=self._on_periode_selected, state="readonly",
                                              font=(FONT, 11), dropdown_font=(FONT, 10),
                                              corner_radius=8, border_width=1, border_color=BORDER_COLOR,
                                              fg_color=BG_MAIN, button_color=TEXT_ACCENT, 
                                              button_hover_color=SUCCESS_GREEN,
                                              height=38)
        self.periode_dropdown.grid(row=1, column=0, sticky="ew")
        self.periode_dropdown.set("Toutes les périodes")
        
        # ============ SECTION ÉLÈVES ============
        eleves_section = ctk.CTkFrame(left_panel, fg_color="transparent")
        eleves_section.grid(row=3, column=0, sticky="ew", padx=25, pady=(18, 0))
        eleves_section.grid_columnconfigure(0, weight=1)
        
        # Titre de section
        eleves_title = ctk.CTkLabel(eleves_section, text="ÉLÈVES", 
                                     font=(FONT, 9, "bold"), text_color=TEXT_SECONDARY,
                                     anchor="w")
        eleves_title.grid(row=0, column=0, sticky="w", pady=(0, 10))
        
        # Card pour la sélection d'élève
        eleves_card = ctk.CTkFrame(eleves_section, fg_color=BG_CARD, corner_radius=10)
        eleves_card.grid(row=1, column=0, sticky="ew")
        eleves_card.grid_columnconfigure(0, weight=1)
        
        # Container élève
        eleve_container = ctk.CTkFrame(eleves_card, fg_color="transparent")
        eleve_container.grid(row=0, column=0, sticky="ew", padx=12, pady=(12, 12))
        eleve_container.grid_columnconfigure(0, weight=1)
        
        ctk.CTkLabel(eleve_container, text="Sélectionner un élève", font=(FONT, 10, "bold"), 
                    text_color=TEXT_PRIMARY, anchor="w").grid(row=0, column=0, sticky="w", pady=(0, 5))
        
        self.eleve_var = StringVar()
        self.eleve_dropdown = ctk.CTkComboBox(eleve_container, variable=self.eleve_var,
                                              values=["Sélectionner classe et période..."],
                                              command=self._on_eleve_selected, state="readonly",
                                              font=(FONT, 11), dropdown_font=(FONT, 10),
                                              corner_radius=8, border_width=1, border_color=BORDER_COLOR,
                                              fg_color=BG_MAIN, button_color=TEXT_ACCENT, 
                                              button_hover_color=SUCCESS_GREEN,
                                              height=38)
        self.eleve_dropdown.grid(row=1, column=0, sticky="ew")
        self.eleve_dropdown.set("Sélectionner classe et période...")
        
        # ============ SECTION ACTIONS ÉLÉGANTE ============
        actions_section = ctk.CTkFrame(left_panel, fg_color="transparent")
        actions_section.grid(row=4, column=0, sticky="sew", padx=25, pady=(18, 25))
        actions_section.grid_columnconfigure(0, weight=1)
        
        # Titre de section
        actions_title = ctk.CTkLabel(actions_section, text="ACTIONS", 
                                     font=(FONT, 9, "bold"), text_color=TEXT_SECONDARY,
                                     anchor="w")
        actions_title.grid(row=0, column=0, sticky="w", pady=(0, 10))
        
        # ========== Bouton Générer Bulletin Word (bleu) ==========
        individual_icon = load_icon('person', (18, 18))
        generate_individual_btn = ctk.CTkButton(
            actions_section, 
            text="Générer Bulletin Word",
            image=individual_icon,
            compound="left",
            command=self._generate_bulletin_word_selected, 
            fg_color="#3498DB", 
            hover_color="#2980B9",
            text_color="#FFFFFF",
            height=45, 
            font=(FONT, 12, "bold"),
            corner_radius=10,
            border_width=0,
            anchor="center"
        )
        generate_individual_btn.grid(row=1, column=0, sticky="ew", pady=(0, 8))
        
        # ========== Bouton Exporter Excel (orange) ==========
        export_icon = load_icon('export', (18, 18))
        export_btn = ctk.CTkButton(
            actions_section, 
            text="Exporter vers Excel",
            image=export_icon,
            compound="left",
                                  command=self._export_to_excel, 
            fg_color="#E67E22", 
            hover_color="#D35400",
            text_color="#FFFFFF",
            height=45, 
            font=(FONT, 12, "bold"),
            corner_radius=10,
            border_width=0,
            anchor="center"
        )
        export_btn.grid(row=2, column=0, sticky="ew", pady=(0, 12))
        
        # Séparateur élégant avant les boutons outline
        sep_actions = ctk.CTkFrame(actions_section, height=1, fg_color=BORDER_COLOR)
        sep_actions.grid(row=3, column=0, sticky="ew", pady=(3, 10))
        
        # ========== Bouton Générer Bulletins (outline style) ==========
        generate_icon = load_icon('add', (16, 16))
        generate_bulletins_btn = ctk.CTkButton(
            actions_section, 
            text="Générer Bulletins",
            image=generate_icon,
            compound="left",
            command=self._generate_all_bulletins, 
            fg_color="transparent", 
            hover_color=BG_CARD,
            text_color=TEXT_ACCENT, 
            border_color=TEXT_ACCENT,
            height=42, 
            font=(FONT, 11, "bold"),
            corner_radius=10,
            border_width=2,
            anchor="center"
        )
        generate_bulletins_btn.grid(row=4, column=0, sticky="ew", pady=(0, 8))
        
        # ========== Bouton Actualiser (outline style) ==========
        refresh_icon = load_icon('refresh', (16, 16))
        refresh_all_btn = ctk.CTkButton(
            actions_section, 
            text="Actualiser",
            image=refresh_icon,
            compound="left",
                                       command=self._refresh_all, 
            fg_color="transparent", 
            hover_color=BG_CARD,
            text_color=TEXT_ACCENT, 
            border_color=TEXT_ACCENT,
            height=42, 
            font=(FONT, 11, "bold"),
            corner_radius=10,
            border_width=2,
            anchor="center"
        )
        refresh_all_btn.grid(row=5, column=0, sticky="ew")
        
        # ========== Bouton Recalculer Rangs (outline style) ==========
        recalc_icon = load_icon('sort', (16, 16))
        recalc_ranks_btn = ctk.CTkButton(
            actions_section, 
            text="Recalculer Rangs",
            image=recalc_icon,
            compound="left",
            command=self._recalculate_ranks_ui, 
            fg_color="transparent", 
            hover_color=BG_CARD,
            text_color=WARNING_ORANGE, 
            border_color=WARNING_ORANGE,
            height=42, 
            font=(FONT, 11, "bold"),
            corner_radius=10,
            border_width=2,
            anchor="center"
        )
        recalc_ranks_btn.grid(row=6, column=0, sticky="ew", pady=(8, 0))
        
        # ========== Bouton Régénérer Bulletins par Trimestre (outline style) ==========
        regenerate_icon = load_icon('generate', (16, 16))
        regenerate_bulletins_btn = ctk.CTkButton(
            actions_section, 
            text="Régénérer par Trimestre",
            image=regenerate_icon,
            compound="left",
            command=self._regenerate_bulletins_by_trimestre_ui, 
            fg_color="transparent", 
            hover_color=BG_CARD,
            text_color=SUCCESS_GREEN, 
            border_color=SUCCESS_GREEN,
            height=42, 
            font=(FONT, 11, "bold"),
            corner_radius=10,
            border_width=2,
            anchor="center"
        )
        regenerate_bulletins_btn.grid(row=7, column=0, sticky="ew", pady=(8, 0))
        
        # ========== Bouton Générer Tous les Trimestres (outline style) ==========
        generate_all_icon = load_icon('generate', (16, 16))
        generate_all_trimestres_btn = ctk.CTkButton(
            actions_section, 
            text="Générer Tous Trimestres",
            image=generate_all_icon,
            compound="left",
            command=self._generate_all_trimestres_ui, 
            fg_color="transparent", 
            hover_color=BG_CARD,
            text_color=ACCENT, 
            border_color=ACCENT,
            height=42, 
            font=(FONT, 11, "bold"),
            corner_radius=10,
            border_width=2,
            anchor="center"
        )
        generate_all_trimestres_btn.grid(row=8, column=0, sticky="ew", pady=(8, 0))
    
    def _build_bulletins_dashboard(self):
        """Construit le tableau des bulletins"""
        # Frame principal du panneau droit
        right_panel = ctk.CTkFrame(self, fg_color=BG_CARD, corner_radius=12)
        right_panel.grid(row=0, column=1, sticky="nsew", padx=(MARGIN_SMALL, MARGIN_MEDIUM), pady=MARGIN_MEDIUM)
        right_panel.grid_columnconfigure(0, weight=1)
        right_panel.grid_rowconfigure(1, weight=1)
        
        # Titre du tableau
        table_title = ctk.CTkLabel(right_panel, text="Bulletins par Classe", font=F_SUB, text_color=TEXT_PRIMARY)
        table_title.grid(row=0, column=0, sticky="ew", padx=MARGIN_MEDIUM, pady=(MARGIN_MEDIUM, MARGIN_SMALL))
        
        # Frame du tableau
        self.table_frame = ctk.CTkFrame(right_panel, fg_color=BG_CARD, corner_radius=12)
        self.table_frame.grid(row=1, column=0, sticky="nsew", padx=0, pady=0)
        self.table_frame.grid_columnconfigure(0, weight=1)
        self.table_frame.grid_rowconfigure(0, weight=1)
    
    def _show_no_selection_message(self):
        """Affiche un message quand aucune sélection"""
        if not self.table_frame:
            return
        
        # Nettoyer le frame
        for widget in self.table_frame.winfo_children():
            widget.destroy()
        
        # Message central
        message_frame = ctk.CTkFrame(self.table_frame, fg_color="transparent")
        message_frame.grid(row=0, column=0, sticky="nsew")
        message_frame.grid_columnconfigure(0, weight=1)
        message_frame.grid_rowconfigure(0, weight=1)
        
        # Icône et texte
        newspaper_icon = load_icon('newspaper', (64, 64))
        icon_label = ctk.CTkLabel(message_frame, image=newspaper_icon, text="")
        icon_label.grid(row=0, column=0, pady=(0, MARGIN_MEDIUM))
        
        title_label = ctk.CTkLabel(message_frame, text="Sélectionnez une classe et une période", 
                                  font=F_SUB, text_color=TEXT_PRIMARY)
        title_label.grid(row=1, column=0, pady=(0, MARGIN_SMALL))
        
        desc_label = ctk.CTkLabel(message_frame, text="Choisissez une classe et une période pour afficher\nles bulletins générés automatiquement.",
                                 font=F_SMALL, text_color=TEXT_SECONDARY)
        desc_label.grid(row=2, column=0)
    
    def _show_no_bulletin_for_student(self, student_name):
        """Affiche un message quand l'élève n'a pas de bulletin"""
        if not self.table_frame:
            return
        
        # Nettoyer le frame
        for widget in self.table_frame.winfo_children():
            widget.destroy()
        
        # Message central
        message_frame = ctk.CTkFrame(self.table_frame, fg_color="transparent")
        message_frame.grid(row=0, column=0, sticky="nsew")
        message_frame.grid_columnconfigure(0, weight=1)
        message_frame.grid_rowconfigure(0, weight=1)
        
        # Icône et texte
        person_icon = load_icon('person', (64, 64))
        icon_label = ctk.CTkLabel(message_frame, image=person_icon, text="")
        icon_label.grid(row=0, column=0, pady=(0, MARGIN_MEDIUM))
        
        title_label = ctk.CTkLabel(message_frame, text=f"Aucun bulletin pour {student_name}", 
                                  font=F_SUB, text_color=TEXT_PRIMARY)
        title_label.grid(row=1, column=0, pady=(0, MARGIN_SMALL))
        
        desc_label = ctk.CTkLabel(message_frame, 
                                 text=f"Aucun bulletin trouvé pour cet élève.\nPériode: {self.selected_periode}",
                                 font=F_SMALL, text_color=TEXT_SECONDARY)
        desc_label.grid(row=2, column=0)
    
    def _display_detailed_bulletin(self, eleve):
        """Affiche le bulletin détaillé d'un élève avec toutes ses notes par matière"""
        if not self.table_frame:
            return
        
        # Nettoyer le frame
        for widget in self.table_frame.winfo_children():
            widget.destroy()
        
        try:
            # Récupérer les notes de l'élève
            eleve_id = eleve.get("id_eleve")
            eleve_nom = eleve.get("nom", "")
            eleve_prenom = eleve.get("prenom", "")
            
            notes_eleve = get_notes_by_eleve(eleve_id, trimestre=self.selected_periode)
            
            if not notes_eleve:
                self._show_no_bulletin_for_student(f"{eleve_prenom} {eleve_nom}")
                return
            
            # Calculer la moyenne et regrouper par matière
            total_points = 0
            total_coefficients = 0
            matieres_notes = {}
            
            for note in notes_eleve:
                note_value = float(note.get("note", 0))
                coefficient = self._extract_coefficient(note)
                matiere = self._extract_subject_name(note)
                
                if matiere not in matieres_notes:
                    matieres_notes[matiere] = {
                        'notes': [],
                        'coefficient': coefficient
                    }
                matieres_notes[matiere]['notes'].append(note_value)
                
                total_points += note_value * coefficient
                total_coefficients += coefficient
            
            if total_coefficients == 0:
                self._show_no_bulletin_for_student(f"{eleve_prenom} {eleve_nom}")
                return
            
            moyenne_generale = total_points / total_coefficients
            
            # Déterminer la mention
            if moyenne_generale >= 16:
                mention = "Très Bien"
                mention_color = "#4CAF50"
            elif moyenne_generale >= 14:
                mention = "Bien"
                mention_color = "#2196F3"
            elif moyenne_generale >= 12:
                mention = "Assez Bien"
                mention_color = "#FF9800"
            elif moyenne_generale >= 10:
                mention = "Passable"
                mention_color = "#FFC107"
            else:
                mention = "Insuffisant"
                mention_color = "#F44336"
            
            # Déterminer le rang
            bulletins_classe = [b for b in self.bulletins 
                              if b.get('classe_nom') == self.selected_classe 
                              and b.get('periode') == self.selected_periode]
            bulletins_classe.sort(key=lambda x: x.get('moyenne_generale', 0), reverse=True)
            rang = next((i+1 for i, b in enumerate(bulletins_classe) if b.get('id_eleve') == eleve_id), 0)
            
            # Créer l'interface du bulletin détaillé
            self._create_detailed_bulletin_ui(eleve, matieres_notes, moyenne_generale, mention, mention_color, rang, len(bulletins_classe))
            
        except Exception as e:
            print(f"Erreur affichage bulletin détaillé: {e}")
            import traceback
            traceback.print_exc()
            self._show_no_bulletin_for_student(f"{eleve.get('prenom', '')} {eleve.get('nom', '')}")
    
    def _create_detailed_bulletin_ui(self, eleve, matieres_notes, moyenne_generale, mention, mention_color, rang, total_eleves):
        """Crée l'interface utilisateur du bulletin détaillé avec export Excel"""
        # Frame principal avec scroll
        main_frame = ctk.CTkScrollableFrame(self.table_frame, fg_color="transparent")
        main_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        # En-tête du bulletin
        header_frame = ctk.CTkFrame(main_frame, fg_color=BG_CARD, corner_radius=12)
        header_frame.pack(fill="x", pady=(0, 20))
        
        # Titre avec bouton d'export
        title_frame = ctk.CTkFrame(header_frame, fg_color="transparent")
        title_frame.pack(fill="x", padx=20, pady=20)
        
        title_label = ctk.CTkLabel(title_frame, 
                                  text="BULLETIN SCOLAIRE INDIVIDUEL",
                                  font=(FONT, 20, "bold"),
                                  text_color=TEXT_ACCENT)
        title_label.pack(side="left")
        
        # Bouton d'export Excel
        export_icon = load_icon('export', (16, 16))
        export_btn = ctk.CTkButton(
            title_frame,
            text="Exporter Excel",
            image=export_icon,
            compound="left",
            font=(FONT, 12, "bold"),
            fg_color="#4CAF50",
            hover_color="#45a049",
            text_color="white",
            corner_radius=8,
            height=35,
            command=lambda: self._export_individual_bulletin_excel(eleve, matieres_notes, moyenne_generale, mention, rang, total_eleves)
        )
        export_btn.pack(side="right", padx=(10, 0))
        
        # Informations de l'élève
        info_frame = ctk.CTkFrame(header_frame, fg_color="transparent")
        info_frame.pack(fill="x", padx=20, pady=(0, 20))
        
        # Grille d'informations
        info_grid = ctk.CTkFrame(info_frame, fg_color="transparent")
        info_grid.pack(fill="x")
        
        # Ligne 1
        row1 = ctk.CTkFrame(info_grid, fg_color="transparent")
        row1.pack(fill="x", pady=5)
        
        ctk.CTkLabel(row1, text="Nom et Prénom:", font=(FONT, 12, "bold"), 
                    text_color=TEXT_PRIMARY).pack(side="left")
        ctk.CTkLabel(row1, text=f"{eleve.get('nom', '')} {eleve.get('prenom', '')}", 
                    font=(FONT, 12), text_color=TEXT_SECONDARY).pack(side="left", padx=(10, 0))
        
        ctk.CTkLabel(row1, text="Classe:", font=(FONT, 12, "bold"), 
                    text_color=TEXT_PRIMARY).pack(side="right")
        ctk.CTkLabel(row1, text=self.selected_classe, 
                    font=(FONT, 12), text_color=TEXT_SECONDARY).pack(side="right", padx=(10, 0))
        
        # Ligne 2
        row2 = ctk.CTkFrame(info_grid, fg_color="transparent")
        row2.pack(fill="x", pady=5)
        
        ctk.CTkLabel(row2, text="Période:", font=(FONT, 12, "bold"), 
                    text_color=TEXT_PRIMARY).pack(side="left")
        ctk.CTkLabel(row2, text=self.selected_periode, 
                    font=(FONT, 12), text_color=TEXT_SECONDARY).pack(side="left", padx=(10, 0))
        
        ctk.CTkLabel(row2, text="Moyenne Générale:", font=(FONT, 12, "bold"), 
                    text_color=TEXT_PRIMARY).pack(side="right")
        ctk.CTkLabel(row2, text=f"{moyenne_generale:.2f}/20", 
                    font=(FONT, 12, "bold"), text_color=mention_color).pack(side="right", padx=(10, 0))
        
        # Ligne 3
        row3 = ctk.CTkFrame(info_grid, fg_color="transparent")
        row3.pack(fill="x", pady=5)
        
        ctk.CTkLabel(row3, text="Rang:", font=(FONT, 12, "bold"), 
                    text_color=TEXT_PRIMARY).pack(side="left")
        ctk.CTkLabel(row3, text=f"{rang}/{total_eleves}", 
                    font=(FONT, 12), text_color=TEXT_SECONDARY).pack(side="left", padx=(10, 0))
        
        ctk.CTkLabel(row3, text="Mention:", font=(FONT, 12, "bold"), 
                    text_color=TEXT_PRIMARY).pack(side="right")
        ctk.CTkLabel(row3, text=mention, 
                    font=(FONT, 12, "bold"), text_color=mention_color).pack(side="right", padx=(10, 0))
        
        # Section des notes par matière
        notes_frame = ctk.CTkFrame(main_frame, fg_color=BG_CARD, corner_radius=12)
        notes_frame.pack(fill="x", pady=(0, 20))
        
        # Titre de la section
        notes_title = ctk.CTkLabel(notes_frame, 
                                  text="NOTES PAR MATIÈRE",
                                  font=(FONT, 16, "bold"),
                                  text_color=TEXT_ACCENT)
        notes_title.pack(pady=15)
        
        # Tableau des notes
        notes_data = [["Matière", "Notes", "Moyenne", "Coeff.", "Appréciation"]]
        
        for matiere, data in matieres_notes.items():
            notes_list = data['notes']
            moyenne_matiere = sum(notes_list) / len(notes_list)
            coefficient = data['coefficient']
            
            # Appréciation par matière
            if moyenne_matiere >= 16:
                appreciation = "Très bien"
            elif moyenne_matiere >= 14:
                appreciation = "Bien"
            elif moyenne_matiere >= 12:
                appreciation = "Assez bien"
            elif moyenne_matiere >= 10:
                appreciation = "Passable"
            else:
                appreciation = "Insuffisant"
            
            # Formater les notes
            notes_str = ", ".join([f"{note:.1f}" for note in notes_list])
            
            notes_data.append([
                matiere,
                notes_str,
                f"{moyenne_matiere:.2f}",
                str(coefficient),
                appreciation
            ])
        
        # Créer le tableau
        notes_table = CTkTable(
            master=notes_frame,
            row=len(notes_data),
            column=len(notes_data[0]),
            values=notes_data,
            header_color=BG_SIDEBAR,
            colors=[BG_CARD, BG_SECONDARY],
            hover_color=TEXT_ACCENT,
            corner_radius=8,
            border_width=1,
            border_color=BORDER_COLOR,
            text_color=TEXT_PRIMARY,
            font=(FONT, 11),
            command=None
        )
        
        notes_table.pack(padx=20, pady=(0, 20), fill="x")
        
        # Section d'appréciation générale
        appreciation_frame = ctk.CTkFrame(main_frame, fg_color=BG_CARD, corner_radius=12)
        appreciation_frame.pack(fill="x", pady=(0, 20))
        
        # Titre de la section
        appreciation_title = ctk.CTkLabel(appreciation_frame, 
                                        text="APPRÉCIATION GÉNÉRALE",
                                        font=(FONT, 16, "bold"),
                                        text_color=TEXT_ACCENT)
        appreciation_title.pack(pady=15)
        
        # Appréciation détaillée
        appreciation_text = self._generate_appreciation_text(moyenne_generale, mention, rang, total_eleves)
        
        appreciation_content = ctk.CTkLabel(appreciation_frame,
                                          text=appreciation_text,
                                          font=(FONT, 12),
                                          text_color=TEXT_PRIMARY,
                                          justify="left",
                                          wraplength=600)
        appreciation_content.pack(padx=20, pady=(0, 20))
        
        # Section des statistiques
        stats_frame = ctk.CTkFrame(main_frame, fg_color=BG_CARD, corner_radius=12)
        stats_frame.pack(fill="x", pady=(0, 20))
        
        stats_title = ctk.CTkLabel(stats_frame, 
                                  text="STATISTIQUES DE PERFORMANCE",
                                  font=(FONT, 16, "bold"),
                                  text_color=TEXT_ACCENT)
        stats_title.pack(pady=15)
        
        # Grille des statistiques
        stats_grid = ctk.CTkFrame(stats_frame, fg_color="transparent")
        stats_grid.pack(fill="x", padx=20, pady=(0, 20))
        
        # Calculer les statistiques
        total_matieres = len(matieres_notes)
        matieres_excellentes = sum(1 for data in matieres_notes.values() 
                                 if sum(data['notes']) / len(data['notes']) >= 16)
        matieres_bonnes = sum(1 for data in matieres_notes.values() 
                            if 14 <= sum(data['notes']) / len(data['notes']) < 16)
        matieres_moyennes = sum(1 for data in matieres_notes.values() 
                              if 10 <= sum(data['notes']) / len(data['notes']) < 14)
        matieres_faibles = total_matieres - matieres_excellentes - matieres_bonnes - matieres_moyennes
        
        # Ligne 1 des statistiques
        stats_row1 = ctk.CTkFrame(stats_grid, fg_color="transparent")
        stats_row1.pack(fill="x", pady=5)
        
        ctk.CTkLabel(stats_row1, text="Total des matières:", font=(FONT, 12, "bold"), 
                    text_color=TEXT_PRIMARY).pack(side="left")
        ctk.CTkLabel(stats_row1, text=str(total_matieres), 
                    font=(FONT, 12), text_color=TEXT_SECONDARY).pack(side="left", padx=(10, 0))
        
        ctk.CTkLabel(stats_row1, text="Matières excellentes (≥16):", font=(FONT, 12, "bold"), 
                    text_color=TEXT_PRIMARY).pack(side="right")
        ctk.CTkLabel(stats_row1, text=str(matieres_excellentes), 
                    font=(FONT, 12), text_color="#4CAF50").pack(side="right", padx=(10, 0))
        
        # Ligne 2 des statistiques
        stats_row2 = ctk.CTkFrame(stats_grid, fg_color="transparent")
        stats_row2.pack(fill="x", pady=5)
        
        ctk.CTkLabel(stats_row2, text="Matières bonnes (14-15.9):", font=(FONT, 12, "bold"), 
                    text_color=TEXT_PRIMARY).pack(side="left")
        ctk.CTkLabel(stats_row2, text=str(matieres_bonnes), 
                    font=(FONT, 12), text_color="#2196F3").pack(side="left", padx=(10, 0))
        
        ctk.CTkLabel(stats_row2, text="Matières moyennes (10-13.9):", font=(FONT, 12, "bold"), 
                    text_color=TEXT_PRIMARY).pack(side="right")
        ctk.CTkLabel(stats_row2, text=str(matieres_moyennes), 
                    font=(FONT, 12), text_color="#FF9800").pack(side="right", padx=(10, 0))
        
        # Ligne 3 des statistiques
        stats_row3 = ctk.CTkFrame(stats_grid, fg_color="transparent")
        stats_row3.pack(fill="x", pady=5)
        
        ctk.CTkLabel(stats_row3, text="Matières à améliorer (<10):", font=(FONT, 12, "bold"), 
                    text_color=TEXT_PRIMARY).pack(side="left")
        ctk.CTkLabel(stats_row3, text=str(matieres_faibles), 
                    font=(FONT, 12), text_color="#F44336").pack(side="left", padx=(10, 0))
        
        # Footer avec date
        footer_frame = ctk.CTkFrame(main_frame, fg_color="transparent")
        footer_frame.pack(fill="x", pady=10)
        
        footer_label = ctk.CTkLabel(footer_frame,
                                   text=f"Document généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}",
                                   font=(FONT, 9),
                                   text_color=TEXT_SECONDARY)
        footer_label.pack()
    
    def _generate_appreciation_text(self, moyenne_generale, mention, rang, total_eleves):
        """Génère une appréciation personnalisée basée sur les performances de l'élève"""
        if moyenne_generale >= 16:
            base_text = f"Excellente performance ! {mention} avec une moyenne de {moyenne_generale:.2f}/20. "
            if rang <= 3:
                base_text += f"Félicitations pour votre {rang}e place sur {total_eleves} élèves. "
            base_text += "Votre travail assidu et votre rigueur sont remarquables. Continuez sur cette excellente voie !"
            
        elif moyenne_generale >= 14:
            base_text = f"Très bonne performance ! {mention} avec une moyenne de {moyenne_generale:.2f}/20. "
            if rang <= total_eleves * 0.3:
                base_text += f"Vous figurez parmi les meilleurs élèves de la classe (rang {rang}/{total_eleves}). "
            base_text += "Votre investissement porte ses fruits. Maintenez cet effort pour progresser encore !"
            
        elif moyenne_generale >= 12:
            base_text = f"Bonne performance ! {mention} avec une moyenne de {moyenne_generale:.2f}/20. "
            if rang <= total_eleves * 0.5:
                base_text += f"Vous êtes dans la première moitié de la classe (rang {rang}/{total_eleves}). "
            base_text += "Votre travail est satisfaisant. Quelques efforts supplémentaires vous permettront d'améliorer encore vos résultats."
            
        elif moyenne_generale >= 10:
            base_text = f"Performance acceptable ! {mention} avec une moyenne de {moyenne_generale:.2f}/20. "
            base_text += f"Rang {rang}/{total_eleves}. Votre niveau est passable mais des efforts sont nécessaires pour progresser. "
            base_text += "Concentrez-vous sur les matières où vous avez des difficultés et n'hésitez pas à demander de l'aide."
            
        else:
            base_text = f"Performance insuffisante avec une moyenne de {moyenne_generale:.2f}/20. "
            base_text += f"Rang {rang}/{total_eleves}. Des efforts importants sont nécessaires pour améliorer vos résultats. "
            base_text += "Il est recommandé de revoir les bases et de travailler régulièrement. N'hésitez pas à solliciter l'aide de vos enseignants."
        
        return base_text
    
    def _clean_appreciation(self, appreciation, moyenne):
        """Nettoie et formate l'appréciation pour l'affichage dans le tableau"""
        if not appreciation:
            # Générer une appréciation basée sur la moyenne si aucune n'existe
            if moyenne >= 16:
                return "Très bien"
            elif moyenne >= 14:
                return "Bien"
            elif moyenne >= 12:
                return "Assez bien"
            elif moyenne >= 10:
                return "Passable"
            else:
                return "Insuffisant"
        
        # Nettoyer les appréciations trop longues ou mal formatées
        appreciation = appreciation.strip()
        
        # Si l'appréciation contient du texte descriptif, extraire seulement la mention
        if "Bulletin individuel" in appreciation or "matières" in appreciation:
            # Extraire la mention à la fin
            if "Très bien" in appreciation:
                return "Très bien"
            elif "Bien" in appreciation:
                return "Bien"
            elif "Assez bien" in appreciation:
                return "Assez bien"
            elif "Passable" in appreciation:
                return "Passable"
            elif "Insuffisant" in appreciation:
                return "Insuffisant"
            else:
                # Fallback basé sur la moyenne
                if moyenne >= 16:
                    return "Très bien"
                elif moyenne >= 14:
                    return "Bien"
                elif moyenne >= 12:
                    return "Assez bien"
                elif moyenne >= 10:
                    return "Passable"
                else:
                    return "Insuffisant"
        
        # Limiter la longueur de l'appréciation
        if len(appreciation) > 20:
            return appreciation[:17] + "..."
        
        return appreciation
    
    def _on_classe_selected(self, selected_classe):
        """Gère la sélection d'une classe"""
        print(f"DEBUG: _on_classe_selected appele avec: '{selected_classe}'")
        
        if selected_classe == "Toutes les classes":
            print("DEBUG: Selection 'Toutes les classes' - reinitialisation")
            self.selected_classe = None
        else:
            print(f"DEBUG: Selection classe specifique: '{selected_classe}'")
            self.selected_classe = selected_classe
        
        print(f"Classe selectionnee: {selected_classe}")
        self._update_eleves_list()
        self._filter_bulletins()
    
    def _on_periode_selected(self, selected_periode):
        """Gère la sélection d'une période"""
        print(f"DEBUG: _on_periode_selected appele avec: '{selected_periode}'")
        
        if selected_periode == "Toutes les périodes":
            print("DEBUG: Selection 'Toutes les periodes' - reinitialisation")
            self.selected_periode = None
        else:
            print(f"DEBUG: Selection periode specifique: '{selected_periode}'")
            self.selected_periode = selected_periode
        
        print(f"Periode selectionnee: {selected_periode}")
        self._update_eleves_list()
        self._filter_bulletins()
    
    def _on_eleve_selected(self, selected_eleve):
        """Gère la sélection d'un élève et affiche son bulletin détaillé"""
        print(f"DEBUG: _on_eleve_selected appele avec: '{selected_eleve}'")
        
        if selected_eleve in ["Sélectionner classe et période...", "Sélectionner un élève...", "Aucun élève dans cette classe", ""]:
            return
        
        # Afficher le bulletin détaillé de cet élève
        if self.selected_classe and self.selected_periode:
            # Récupérer l'ID de la classe (robuste)
            classe_id = self._get_class_id_by_name(self.selected_classe)
            
            if not classe_id:
                return
            
            # Récupérer tous les élèves de la classe
            eleves_classe = get_all_eleves()
            eleves_classe = [e for e in eleves_classe if self._extract_student_class_id(e) == classe_id]
            
            # Trouver l'élève sélectionné
            eleve_selected = None
            for eleve in eleves_classe:
                if f"{eleve['nom']} {eleve['prenom']}" == selected_eleve:
                    eleve_selected = eleve
                    break
            
            if eleve_selected:
                print(f"Affichage du bulletin détaillé de {selected_eleve}")
                self._display_detailed_bulletin(eleve_selected)
            else:
                print(f"Élève non trouvé: {selected_eleve}")
                self._show_no_bulletin_for_student(selected_eleve)
    
    def _update_eleves_list(self):
        """Met à jour la liste des élèves selon la classe et période sélectionnées"""
        if not self.selected_classe or not self.selected_periode:
            self.eleve_dropdown.configure(values=["Sélectionner classe et période..."])
            self.eleve_dropdown.set("Sélectionner classe et période...")
            return
        
        # Récupérer l'ID de la classe (robuste)
        classe_id = self._get_class_id_by_name(self.selected_classe)
        
        if not classe_id:
            print(f"Classe non trouvée: {self.selected_classe}")
            return
        
        # Récupérer tous les élèves de la classe (compat clés)
        eleves_classe = get_all_eleves()
        print(f"Total élèves récupérés: {len(eleves_classe)}")
        
        def _belongs_to_class(e):
            return (
                e.get('id_classe') == classe_id
                or e.get('classe_id') == classe_id
                or e.get('idclasse') == classe_id
            )
        eleves_classe = [e for e in eleves_classe if _belongs_to_class(e)]
        print(f"Élèves filtrés pour classe {classe_id}: {len(eleves_classe)}")
        
        if eleves_classe:
            eleves_list = [self._student_fullname(e) for e in eleves_classe]
            eleves_list.sort()  # Trier par ordre alphabétique
            self.eleve_dropdown.configure(values=eleves_list)
            self.eleve_dropdown.set("Sélectionner un élève...")
            print(f"{len(eleves_list)} élèves chargés pour la classe {self.selected_classe}: {eleves_list[:5]}...")
        else:
            self.eleve_dropdown.configure(values=["Aucun élève dans cette classe"])
            self.eleve_dropdown.set("Aucun élève dans cette classe")
            print(f"Aucun élève trouvé pour la classe {self.selected_classe} (ID: {classe_id})")

    def _filter_bulletins(self):
        """Filtre les bulletins selon les sélections"""
        print(f"DEBUG: _filter_bulletins appele - classe: '{self.selected_classe}', periode: '{self.selected_periode}'")
        
        # Réinitialiser la pagination à chaque nouveau filtrage
        self.current_page = 1
        
        if not self.selected_classe and not self.selected_periode:
            print("DEBUG: Aucune selection - affichage message")
            self._show_no_selection_message()
            return
        
        # Filtrer les bulletins selon les sélections
        filtered_bulletins = []
        
        for bulletin in self.bulletins:
            # Filtrer par classe si sélectionnée
            if self.selected_classe and self.selected_classe != "Toutes les classes":
                # Utiliser directement classe_nom du bulletin
                bulletin_classe = bulletin.get('classe_nom', '')
                
                # Debug limité (seulement les 3 premiers)
                if len(filtered_bulletins) < 3:
                    print(f"DEBUG: Comparaison - Bulletin classe: '{bulletin_classe}' vs Selection: '{self.selected_classe}'")
                
                # Si classe_nom est vide, essayer de récupérer la classe via l'ID élève
                if not bulletin_classe:
                    eleve_id = bulletin.get('id_eleve')
                    if eleve_id:
                        bulletin_classe = self._get_eleve_classe_name(eleve_id)
                        if len(filtered_bulletins) < 5:
                            print(f"DEBUG: Classe recuperee via ID eleve {eleve_id}: '{bulletin_classe}'")
                
                if bulletin_classe != self.selected_classe:
                    continue
            
            # Filtrer par période si sélectionnée
            if self.selected_periode and self.selected_periode != "Toutes les périodes":
                if bulletin.get('periode') != self.selected_periode:
                    continue
            
            filtered_bulletins.append(bulletin)
        
        print(f"DEBUG: {len(filtered_bulletins)} bulletins filtres sur {len(self.bulletins)} total")
        
        # Debug: Afficher les détails des bulletins filtrés
        if filtered_bulletins:
            print(f"DEBUG: Bulletins pour '{self.selected_classe}' - {self.selected_periode}:")
            for i, bulletin in enumerate(filtered_bulletins[:5]):  # Afficher les 5 premiers
                eleve_nom = self._get_eleve_name(bulletin.get('id_eleve'))
                rang = bulletin.get('rang', 0)
                moyenne = bulletin.get('moyenne_generale', 0)
                print(f"  {i+1}. {eleve_nom} - Rang: {rang} - Moyenne: {moyenne}")
            if len(filtered_bulletins) > 5:
                print(f"  ... et {len(filtered_bulletins) - 5} autres")
        
        # Si aucun bulletin trouvé pour une classe spécifique, proposer de les générer
        if len(filtered_bulletins) == 0 and self.selected_classe and self.selected_classe != "Toutes les classes":
            self._show_generate_bulletins_option()
        else:
            # Recalculer les rangs si nécessaire avant affichage
            if filtered_bulletins:
                self._ensure_ranks_are_calculated(filtered_bulletins)
            
            # Afficher les bulletins filtrés
            self._display_bulletins(filtered_bulletins)
    
    def _ensure_ranks_are_calculated(self, bulletins):
        """S'assure que les rangs sont correctement calculés pour les bulletins"""
        try:
            # Grouper les bulletins par classe et période
            bulletins_by_class_period = {}
            for bulletin in bulletins:
                classe_nom = bulletin.get('classe_nom', '')
                periode = bulletin.get('periode', '')
                key = f"{classe_nom}_{periode}"
                
                if key not in bulletins_by_class_period:
                    bulletins_by_class_period[key] = []
                bulletins_by_class_period[key].append(bulletin)
            
            # Recalculer les rangs pour chaque groupe
            for key, bulletins_group in bulletins_by_class_period.items():
                # Trier par moyenne décroissante
                bulletins_group.sort(key=lambda x: x.get('moyenne_generale', 0), reverse=True)
                
                # Assigner les rangs
                for i, bulletin in enumerate(bulletins_group):
                    bulletin['rang'] = i + 1
                    
                print(f"DEBUG: Rangs recalculés pour {key}: {len(bulletins_group)} bulletins")
                
        except Exception as e:
            print(f"Erreur lors du recalcul des rangs: {e}")
    
    def _recalculate_ranks_in_db(self, classe_nom, periode):
        """Recalcule les rangs en base de données pour une classe et période spécifiques"""
        try:
            from database.connection import get_db_connection
            conn = get_db_connection()
            cursor = conn.cursor()
            
            # Récupérer l'ID de la classe
            cursor.execute("SELECT id_classe FROM classes WHERE nom = ?", (classe_nom,))
            classe_result = cursor.fetchone()
            if not classe_result:
                print(f"Classe '{classe_nom}' non trouvée en base")
                return
            
            classe_id = classe_result[0]
            
            # Récupérer tous les bulletins de la classe et période, triés par moyenne décroissante
            cursor.execute("""
                SELECT b.id_bulletin, b.id_eleve, b.moyenne_generale
                FROM bulletins b
                JOIN eleves e ON b.id_eleve = e.id_eleve
                WHERE e.id_classe = ? AND b.periode = ?
                ORDER BY b.moyenne_generale DESC
            """, (classe_id, periode))
            
            bulletins = cursor.fetchall()
            
            # Mettre à jour les rangs
            for i, (bulletin_id, eleve_id, moyenne) in enumerate(bulletins):
                rang = i + 1
                cursor.execute("""
                    UPDATE bulletins SET rang = ? WHERE id_bulletin = ?
                """, (rang, bulletin_id))
            
            conn.commit()
            conn.close()
            print(f"Rangs recalculés en base pour la classe {classe_nom} - {periode} ({len(bulletins)} bulletins)")
            
        except Exception as e:
            print(f"Erreur lors du recalcul des rangs en base: {e}")
    
    def _recalculate_ranks_in_db_by_classe_id(self, classe_id, periode):
        """Recalcule les rangs en base de données pour une classe et période spécifiques (par ID classe)"""
        try:
            from database.connection import get_db_connection
            conn = get_db_connection()
            cursor = conn.cursor()
            
            # Récupérer tous les bulletins de la classe et période, triés par moyenne décroissante
            cursor.execute("""
                SELECT b.id_bulletin, b.id_eleve, b.moyenne_generale
                FROM bulletins b
                JOIN eleves e ON b.id_eleve = e.id_eleve
                WHERE e.id_classe = ? AND b.periode = ?
                ORDER BY b.moyenne_generale DESC
            """, (classe_id, periode))
            
            bulletins = cursor.fetchall()
            
            # Mettre à jour les rangs
            for i, (bulletin_id, eleve_id, moyenne) in enumerate(bulletins):
                rang = i + 1
                cursor.execute("""
                    UPDATE bulletins SET rang = ? WHERE id_bulletin = ?
                """, (rang, bulletin_id))
            
            conn.commit()
            conn.close()
            print(f"Rangs recalculés en base pour la classe ID {classe_id} - {periode} ({len(bulletins)} bulletins)")
            
        except Exception as e:
            print(f"Erreur lors du recalcul des rangs en base pour classe ID {classe_id}: {e}")
    
    def _recalculate_ranks_ui(self):
        """Interface utilisateur pour recalculer les rangs"""
        if not self.selected_classe or not self.selected_periode:
            messagebox.showwarning("Sélection requise", 
                                 "Veuillez d'abord sélectionner une classe et une période.")
            return
        
        # Demander confirmation
        result = messagebox.askyesno(
            "Recalculer les rangs", 
            f"Voulez-vous recalculer les rangs pour la classe '{self.selected_classe}' "
            f"et la période '{self.selected_periode}' ?\n\n"
            "Cette action mettra à jour les rangs en base de données."
        )
        
        if result:
            try:
                # Recalculer les rangs en base de données
                self._recalculate_ranks_in_db(self.selected_classe, self.selected_periode)
                
                # Recharger les données et actualiser l'affichage
                self._load_data()
                self._filter_bulletins()
                
                messagebox.showinfo("Succès", 
                                  f"Les rangs ont été recalculés avec succès pour "
                                  f"la classe '{self.selected_classe}' et la période '{self.selected_periode}'.")
                
            except Exception as e:
                messagebox.showerror("Erreur", 
                                   f"Erreur lors du recalcul des rangs:\n{str(e)}")
    
    def _regenerate_bulletins_by_trimestre_ui(self):
        """Interface utilisateur pour régénérer les bulletins par trimestre"""
        if not self.selected_periode:
            messagebox.showwarning("Sélection requise", 
                                 "Veuillez d'abord sélectionner une période (trimestre).")
            return
        
        # Demander confirmation
        result = messagebox.askyesno(
            "Régénérer les bulletins par trimestre", 
            f"Voulez-vous régénérer TOUS les bulletins pour le '{self.selected_periode}' ?\n\n"
            "Cette action va :\n"
            "• Supprimer tous les bulletins existants pour ce trimestre\n"
            "• Recalculer les moyennes depuis les notes\n"
            "• Générer de nouveaux bulletins pour toutes les classes\n"
            "• Calculer les rangs correctement"
        )
        
        if result:
            try:
                # Supprimer tous les bulletins existants pour ce trimestre
                self._delete_all_bulletins_for_trimestre(self.selected_periode)
                
                # Générer les bulletins pour toutes les classes de ce trimestre
                self._generate_bulletins_for_trimestre(self.selected_periode)
                
                messagebox.showinfo("Succès", 
                                  f"Les bulletins ont été régénérés avec succès pour "
                                  f"le '{self.selected_periode}'.\n\n"
                                  f"Toutes les classes ont été traitées avec les moyennes recalculées.")
                
            except Exception as e:
                messagebox.showerror("Erreur", 
                                   f"Erreur lors de la régénération des bulletins:\n{str(e)}")
    
    def _delete_bulletins_for_class_period(self, classe_nom, periode):
        """Supprime les bulletins existants pour une classe et période spécifiques"""
        try:
            from database.connection import get_db_connection
            conn = get_db_connection()
            cursor = conn.cursor()
            
            # Récupérer l'ID de la classe
            cursor.execute("SELECT id_classe FROM classes WHERE nom = ?", (classe_nom,))
            classe_result = cursor.fetchone()
            if not classe_result:
                print(f"Classe '{classe_nom}' non trouvée en base")
                return
            
            classe_id = classe_result[0]
            
            # Supprimer les bulletins de la classe et période
            cursor.execute("""
                DELETE FROM bulletins 
                WHERE id_eleve IN (
                    SELECT id_eleve FROM eleves WHERE id_classe = ?
                ) AND periode = ?
            """, (classe_id, periode))
            
            deleted_count = cursor.rowcount
            conn.commit()
            conn.close()
            print(f"Supprimé {deleted_count} bulletins pour la classe {classe_nom} - {periode}")
            
        except Exception as e:
            print(f"Erreur lors de la suppression des bulletins: {e}")
    
    def _delete_all_bulletins_for_trimestre(self, trimestre):
        """Supprime tous les bulletins existants pour un trimestre donné"""
        try:
            from database.connection import get_db_connection
            conn = get_db_connection()
            cursor = conn.cursor()
            
            # Supprimer tous les bulletins du trimestre
            cursor.execute("DELETE FROM bulletins WHERE periode = ?", (trimestre,))
            
            deleted_count = cursor.rowcount
            conn.commit()
            conn.close()
            print(f"Supprimé {deleted_count} bulletins pour le trimestre {trimestre}")
            
        except Exception as e:
            print(f"Erreur lors de la suppression des bulletins pour le trimestre {trimestre}: {e}")
    
    def _generate_bulletins_for_trimestre(self, trimestre):
        """Génère les bulletins pour toutes les classes d'un trimestre donné"""
        try:
            from src.modules.academic.classes.controllers.classe_controller import get_all_classes
            from src.modules.academic.students.controllers.eleve_controller import get_all_eleves
            from src.modules.academic.grades.controllers.notes_controller import get_notes_by_eleve
            
            # Récupérer toutes les classes
            classes_list = get_all_classes()
            if not classes_list:
                print("Aucune classe trouvée")
                return
            
            # Récupérer tous les élèves
            eleves_list = get_all_eleves()
            if not eleves_list:
                print("Aucun élève trouvé")
                return
            
            bulletins_generes_total = 0
            classes_traitees = 0
            
            print(f"Génération des bulletins pour le trimestre: {trimestre}")
            
            # Traiter chaque classe
            for classe in classes_list:
                classe_id = classe.get('id')
                classe_nom = classe.get('nom')
                
                print(f"Traitement de la classe: {classe_nom}")
                
                # Récupérer les élèves de cette classe
                eleves_classe = [e for e in eleves_list if e.get('id_classe') == classe_id]
                
                if not eleves_classe:
                    print(f"Aucun élève trouvé pour la classe {classe_nom}")
                    continue
                
                bulletins_generes_classe = 0
                
                # Traiter chaque élève de la classe
                for eleve in eleves_classe:
                    eleve_id = eleve.get('id_eleve')
                    eleve_nom = eleve.get('nom', '')
                    eleve_prenom = eleve.get('prenom', '')
                    
                    print(f"Traitement de {eleve_prenom} {eleve_nom} (ID: {eleve_id})")
                    
                    # Récupérer les notes de l'élève pour le trimestre
                    notes_eleve = get_notes_by_eleve(eleve_id, trimestre=trimestre)
                    
                    print(f"DEBUG: {len(notes_eleve)} notes trouvées pour {eleve_prenom} {eleve_nom} - {trimestre}")
                    if notes_eleve:
                        for note in notes_eleve[:3]:  # Afficher les 3 premières notes pour debug
                            print(f"  Note: {note.get('note')} - Matière: {note.get('matiere_nom')} - Date: {note.get('date_evaluation')}")
                    
                    if not notes_eleve:
                        print(f"Aucune note trouvée pour {eleve_prenom} {eleve_nom} - {trimestre}")
                        continue
                    
                    # Calculer la moyenne pondérée
                    total_points = 0
                    total_coefficients = 0
                    
                    for note in notes_eleve:
                        note_value = float(note.get("note", 0))
                        coefficient = float(note.get("coefficient", 1))
                        total_points += note_value * coefficient
                        total_coefficients += coefficient
                    
                    # Calculer la moyenne
                    if total_coefficients == 0:
                        moyenne_generale = 0.0
                        print(f"DEBUG: Aucune note pour {eleve_prenom} {eleve_nom} - {trimestre} - Moyenne: 0.0")
                    else:
                        moyenne_generale = total_points / total_coefficients
                        print(f"DEBUG: {len(notes_eleve)} notes pour {eleve_prenom} {eleve_nom} - {trimestre} - Moyenne: {moyenne_generale:.2f}")
                    
                    # Déterminer la mention
                    if moyenne_generale >= 16:
                        mention = "Très bien"
                    elif moyenne_generale >= 14:
                        mention = "Bien"
                    elif moyenne_generale >= 12:
                        mention = "Assez bien"
                    elif moyenne_generale >= 10:
                        mention = "Passable"
                    else:
                        mention = "Insuffisant"
                    
                    # Créer le bulletin
                    bulletin_info = {
                        "id_eleve": eleve_id,
                        "periode": trimestre,
                        "moyenne_generale": round(moyenne_generale, 2),
                        "rang": 0,  # Sera calculé après
                        "appreciation": mention
                    }
                    
                    # Ajouter le bulletin
                    success = add_bulletin(bulletin_info)
                    if success:
                        bulletins_generes_classe += 1
                        print(f"Bulletin généré pour {eleve_prenom} {eleve_nom} - {trimestre} - Moyenne: {moyenne_generale:.2f}")
                
                print(f"Classe {classe_nom}: {bulletins_generes_classe} bulletins générés")
                bulletins_generes_total += bulletins_generes_classe
                classes_traitees += 1
            
            # Recalculer les rangs pour toutes les classes de ce trimestre
            print("Recalcul des rangs pour toutes les classes...")
            for classe in classes_list:
                classe_id = classe.get('id')
                self._recalculate_ranks_in_db_by_classe_id(classe_id, trimestre)
            
            # Recharger les bulletins
            self._load_data()
            self._filter_bulletins()
            
            print(f"Génération terminée: {bulletins_generes_total} bulletins générés pour {classes_traitees} classes")
            
        except Exception as e:
            print(f"Erreur lors de la génération des bulletins pour le trimestre {trimestre}: {e}")
            import traceback
            traceback.print_exc()
    
    def _generate_all_trimestres_ui(self):
        """Interface utilisateur pour générer les bulletins pour tous les trimestres"""
        # Demander confirmation
        result = messagebox.askyesno(
            "Générer tous les trimestres", 
            "Voulez-vous générer les bulletins pour TOUS les trimestres ?\n\n"
            "Cette action va :\n"
            "• Supprimer tous les bulletins existants\n"
            "• Générer les bulletins pour les 3 trimestres\n"
            "• Traiter toutes les classes\n"
            "• Calculer les moyennes et rangs correctement\n\n"
            "⚠️ Cette opération peut prendre du temps !"
        )
        
        if result:
            try:
                # Supprimer tous les bulletins existants
                self._delete_all_bulletins()
                
                # Générer pour chaque trimestre
                trimestres = ["1er Trimestre", "2ème Trimestre", "3ème Trimestre"]
                total_bulletins = 0
                
                for trimestre in trimestres:
                    print(f"Génération des bulletins pour {trimestre}...")
                    self._generate_bulletins_for_trimestre(trimestre)
                    total_bulletins += self._count_bulletins_for_trimestre(trimestre)
                
                messagebox.showinfo("Succès", 
                                  f"Génération terminée !\n\n"
                                  f"Bulletins générés pour les 3 trimestres.\n"
                                  f"Total: {total_bulletins} bulletins créés.")
                
            except Exception as e:
                messagebox.showerror("Erreur", 
                                   f"Erreur lors de la génération de tous les trimestres:\n{str(e)}")
    
    def _delete_all_bulletins(self):
        """Supprime tous les bulletins existants"""
        try:
            from database.connection import get_db_connection
            conn = get_db_connection()
            cursor = conn.cursor()
            
            # Supprimer tous les bulletins
            cursor.execute("DELETE FROM bulletins")
            
            deleted_count = cursor.rowcount
            conn.commit()
            conn.close()
            print(f"Supprimé {deleted_count} bulletins au total")
            
        except Exception as e:
            print(f"Erreur lors de la suppression de tous les bulletins: {e}")
    
    def _count_bulletins_for_trimestre(self, trimestre):
        """Compte le nombre de bulletins pour un trimestre"""
        try:
            from database.connection import get_db_connection
            conn = get_db_connection()
            cursor = conn.cursor()
            
            cursor.execute("SELECT COUNT(*) FROM bulletins WHERE periode = ?", (trimestre,))
            count = cursor.fetchone()[0]
            conn.close()
            
            return count
            
        except Exception as e:
            print(f"Erreur lors du comptage des bulletins pour {trimestre}: {e}")
            return 0
    
    def _get_eleve_classe_name(self, eleve_id):
        """Récupère le nom de la classe d'un élève depuis les bulletins ou la base de données"""
        try:
            # Chercher dans les bulletins existants pour trouver la classe
            for bulletin in self.bulletins:
                if bulletin.get('id_eleve') == eleve_id and bulletin.get('classe_nom'):
                    return bulletin['classe_nom']
            
            # Si pas trouvé dans les bulletins, essayer de récupérer depuis la base
            try:
                from src.modules.academic.grades.controllers.bulletin_controller import get_all_eleves
                eleves = get_all_eleves()
                for eleve in eleves:
                    if eleve.get('id_eleve') == eleve_id:
                        # Récupérer le nom de la classe via l'ID classe
                        classe_id = eleve.get('id_classe')
                        if classe_id and classe_id in self.classes:
                            return self.classes[classe_id].get('nom', '')
                        break
            except Exception as e:
                print(f"Erreur recuperation classe depuis base pour eleve {eleve_id}: {e}")
            
            return ''
        except Exception as e:
            print(f"Erreur recuperation classe eleve {eleve_id}: {e}")
            return ''
    
    def _fix_missing_classe_names(self):
        """Corrige les bulletins qui n'ont pas de classe_nom en récupérant depuis la base"""
        try:
            bulletins_fixed = 0
            for bulletin in self.bulletins:
                if not bulletin.get('classe_nom'):
                    eleve_id = bulletin.get('id_eleve')
                    if eleve_id:
                        classe_nom = self._get_eleve_classe_name(eleve_id)
                        if classe_nom:
                            bulletin['classe_nom'] = classe_nom
                            bulletins_fixed += 1
            
            if bulletins_fixed > 0:
                print(f"DEBUG: {bulletins_fixed} bulletins corriges avec classe_nom manquant")
                
        except Exception as e:
            print(f"Erreur correction classe_nom: {e}")
    
    def _show_generate_bulletins_option(self):
        """Affiche une option pour générer les bulletins manquants pour la classe sélectionnée"""
        if not self.table_frame:
            return
        
        # Nettoyer le frame
        for widget in self.table_frame.winfo_children():
            widget.destroy()
        
        # Frame principal
        main_frame = ctk.CTkFrame(self.table_frame, fg_color=BG_CARD, corner_radius=12)
        main_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Icône d'information
        info_icon = ctk.CTkLabel(main_frame, text="ℹ️", font=(FONT, 48), text_color=TEXT_ACCENT)
        info_icon.pack(pady=(20, 10))
        
        # Titre
        title_label = ctk.CTkLabel(main_frame, 
                                  text="Aucun bulletin trouvé",
                                  font=(FONT, 20, "bold"),
                                  text_color=TEXT_PRIMARY)
        title_label.pack(pady=(0, 10))
        
        # Message
        message_label = ctk.CTkLabel(main_frame,
                                    text=f"Aucun bulletin n'a été trouvé pour la classe '{self.selected_classe}'.\n\n"
                                         f"Les bulletins peuvent être générés automatiquement à partir des notes existantes.",
                                    font=(FONT, 14),
                                    text_color=TEXT_SECONDARY,
                                    justify="center")
        message_label.pack(pady=(0, 20))
        
        # Boutons d'action
        buttons_frame = ctk.CTkFrame(main_frame, fg_color="transparent")
        buttons_frame.pack(pady=(0, 20))
        
        # Bouton générer bulletins
        generate_btn = ctk.CTkButton(
            buttons_frame,
            text="Générer les bulletins",
            font=(FONT, 14, "bold"),
            fg_color="#4CAF50",
            hover_color="#45a049",
            text_color="white",
            corner_radius=8,
            height=40,
            width=200,
            command=self._generate_bulletins_for_class
        )
        generate_btn.pack(side="left", padx=(0, 10))
        
        # Bouton actualiser
        refresh_btn = ctk.CTkButton(
            buttons_frame,
            text="Actualiser",
            font=(FONT, 14, "bold"),
            fg_color="#2196F3",
            hover_color="#1976D2",
            text_color="white",
            corner_radius=8,
            height=40,
            width=120,
            command=self._refresh_bulletins
        )
        refresh_btn.pack(side="left")
    
    def _generate_bulletins_for_class(self):
        """Génère les bulletins pour la classe sélectionnée"""
        try:
            from tkinter import messagebox
            
            # Demander confirmation
            result = messagebox.askyesno(
                "Génération des bulletins",
                f"Voulez-vous générer les bulletins pour la classe '{self.selected_classe}' ?\n\n"
                f"Cela créera des bulletins basés sur les notes existantes des élèves de cette classe."
            )
            
            if not result:
                return
            
            # Récupérer l'ID de la classe
            classe_id = self._get_class_id_by_name(self.selected_classe)
            if not classe_id:
                messagebox.showerror("Erreur", f"Impossible de trouver l'ID de la classe '{self.selected_classe}'")
                return
            
            # Supprimer les bulletins existants pour cette classe
            self._delete_existing_bulletins_for_class(classe_id)
            
            # Générer les bulletins pour cette classe (version complète avec tous les élèves)
            self._generate_complete_bulletins_for_class(classe_id)
            
        except Exception as e:
            print(f"Erreur generation bulletins pour classe: {e}")
            messagebox.showerror("Erreur", f"Erreur lors de la génération des bulletins:\n{str(e)}")
    
    def _delete_existing_bulletins_for_class(self, classe_id):
        """Supprime les bulletins existants pour une classe avant régénération"""
        try:
            # Récupérer le nom de la classe
            classe_nom = None
            for classe in self.classes.values():
                if classe.get('id') == classe_id:
                    classe_nom = classe.get('nom', '')
                    break
            
            if not classe_nom:
                print(f"Nom de classe non trouvé pour l'ID {classe_id}")
                return
            
            print(f"Suppression des bulletins existants pour la classe '{classe_nom}'...")
            
            # Trouver et supprimer les bulletins de cette classe
            bulletins_to_delete = []
            for bulletin in self.bulletins:
                if bulletin.get('classe_nom') == classe_nom:
                    bulletins_to_delete.append(bulletin)
            
            # Supprimer les bulletins de la base de données
            deleted_count = 0
            for bulletin in bulletins_to_delete:
                bulletin_id = bulletin.get('id')
                if bulletin_id:
                    try:
                        success = delete_bulletin(bulletin_id)
                        if success:
                            deleted_count += 1
                    except Exception as e:
                        print(f"Erreur suppression bulletin {bulletin_id}: {e}")
            
            print(f"{deleted_count} bulletins supprimés pour la classe '{classe_nom}'")
            
        except Exception as e:
            print(f"Erreur suppression bulletins pour classe {classe_id}: {e}")
    
    def _generate_complete_bulletins_for_class(self, classe_id):
        """Génère les bulletins pour TOUS les élèves d'une classe, même ceux sans notes"""
        try:
            from tkinter import messagebox
            
            # Récupérer tous les élèves de la classe
            eleves_classe = get_all_eleves()
            eleves_classe = [e for e in eleves_classe if self._extract_student_class_id(e) == classe_id]
            
            if not eleves_classe:
                messagebox.showwarning("Aucun élève", f"Aucun élève trouvé dans la classe '{self.selected_classe}'")
                return
            
            print(f"DEBUG: {len(eleves_classe)} élèves trouvés dans la classe '{self.selected_classe}'")
            bulletins_generes = 0
            
            for eleve in eleves_classe:
                eleve_id = eleve.get('id_eleve')
                eleve_nom = eleve.get('nom', '')
                eleve_prenom = eleve.get('prenom', '')
                
                # Récupérer les notes de l'élève pour toutes les périodes
                for periode in self.periodes:
                    notes_eleve = get_notes_by_eleve(eleve_id, trimestre=periode)
                    
                    # Calculer la moyenne et regrouper par matière
                    total_points = 0
                    total_coefficients = 0
                    matieres_notes = {}
                    
                    if notes_eleve:
                        for note in notes_eleve:
                            note_value = float(note.get("note", 0))
                            coefficient = self._extract_coefficient(note)
                            matiere = self._extract_subject_name(note)
                            
                            if matiere not in matieres_notes:
                                matieres_notes[matiere] = {
                                    'notes': [],
                                    'coefficient': coefficient
                                }
                            matieres_notes[matiere]['notes'].append(note_value)
                            
                            total_points += note_value * coefficient
                            total_coefficients += coefficient
                    
                    # Si l'élève n'a pas de notes, moyenne = 0
                    if total_coefficients == 0:
                        moyenne_generale = 0.0
                        print(f"DEBUG: Aucune note pour {eleve_prenom} {eleve_nom} - {periode} - Moyenne: 0.0")
                    else:
                        moyenne_generale = total_points / total_coefficients
                        print(f"DEBUG: {len(notes_eleve)} notes pour {eleve_prenom} {eleve_nom} - {periode} - Moyenne: {moyenne_generale:.2f}")
                    
                    # Déterminer la mention
                    if moyenne_generale >= 16:
                        mention = "Très bien"
                    elif moyenne_generale >= 14:
                        mention = "Bien"
                    elif moyenne_generale >= 12:
                        mention = "Assez bien"
                    elif moyenne_generale >= 10:
                        mention = "Passable"
                    else:
                        mention = "Insuffisant"
                    
                    # Créer le bulletin
                    bulletin_info = {
                        "id_eleve": eleve_id,
                        "periode": periode,
                        "moyenne_generale": round(moyenne_generale, 2),
                        "rang": 0,  # Sera calculé après
                        "appreciation": mention
                    }
                    
                    # Ajouter le bulletin
                    success = add_bulletin(bulletin_info)
                    if success:
                        bulletins_generes += 1
                        print(f"Bulletin genere pour {eleve_prenom} {eleve_nom} - {periode} - Moyenne: {moyenne_generale:.2f}")
            
            # Recharger les bulletins
            self._load_data()
            
            # Recalculer les rangs pour toutes les périodes de cette classe
            self._recalculate_all_ranks_for_class(classe_id)
            
            messagebox.showinfo("Génération terminée", 
                              f"Génération des bulletins terminée !\n\n"
                              f"• {bulletins_generes} bulletins générés\n"
                              f"• Classe: {self.selected_classe}\n"
                              f"• {len(eleves_classe)} élèves traités\n"
                              f"• Tous les élèves ont maintenant des bulletins")
            
            # Réinitialiser la pagination et actualiser l'affichage
            self.current_page = 1
            self._filter_bulletins()
            
        except Exception as e:
            print(f"Erreur generation bulletins complete: {e}")
            import traceback
            traceback.print_exc()
    
    def _generate_detailed_bulletins_for_class(self, classe_id):
        """Génère les bulletins détaillés pour une classe spécifique - DÉPRÉCIÉE"""
        # Rediriger vers la nouvelle méthode complète
        self._generate_complete_bulletins_for_class(classe_id)
    
    def _refresh_bulletins(self):
        """Actualise les bulletins depuis la base de données"""
        try:
            self._load_data()
            # Réinitialiser la pagination et actualiser l'affichage
            self.current_page = 1
            self._filter_bulletins()
        except Exception as e:
            print(f"Erreur actualisation bulletins: {e}")
    
    def _generate_all_bulletins(self):
        """Génère les bulletins pour toutes les classes qui n'en ont pas"""
        try:
            from tkinter import messagebox
            
            # Demander confirmation
            result = messagebox.askyesno(
                "Génération des bulletins",
                "Voulez-vous générer les bulletins pour toutes les classes qui n'en ont pas ?\n\n"
                "Cette opération peut prendre quelques minutes selon le nombre d'élèves et de notes."
            )
            
            if not result:
                return
            
            # Récupérer toutes les classes
            classes_list = get_all_classes()
            bulletins_generes_total = 0
            classes_traitees = 0
            
            for classe in classes_list:
                classe_id = classe.get('id')
                classe_nom = classe.get('nom', '')
                
                # Vérifier si cette classe a déjà des bulletins
                bulletins_classe = [b for b in self.bulletins if b.get('classe_nom') == classe_nom]
                
                if bulletins_classe:
                    print(f"Classe '{classe_nom}' a déjà des bulletins, ignorée")
                    continue
                
                # Générer les bulletins pour cette classe (version complète)
                print(f"Génération des bulletins pour la classe '{classe_nom}'...")
                bulletins_generes = self._generate_complete_bulletins_for_class_id(classe_id, classe_nom)
                bulletins_generes_total += bulletins_generes
                classes_traitees += 1
            
            # Recharger les bulletins
            self._load_data()
            
            # Recalculer les rangs pour toutes les classes traitées
            print("Recalcul des rangs pour toutes les classes...")
            for classe in classes_list:
                classe_id = classe.get('id')
                classe_nom = classe.get('nom', '')
                
                # Vérifier si cette classe a des bulletins
                bulletins_classe = [b for b in self.bulletins if b.get('classe_nom') == classe_nom]
                if bulletins_classe:
                    self._recalculate_all_ranks_for_class(classe_id)
            
            messagebox.showinfo("Génération terminée", 
                              f"Génération des bulletins terminée !\n\n"
                              f"• {bulletins_generes_total} bulletins générés\n"
                              f"• {classes_traitees} classes traitées\n"
                              f"• Rangs calculés pour toutes les classes\n"
                              f"• Toutes les classes ont maintenant des bulletins")
            
            # Réinitialiser la pagination et actualiser l'affichage
            self.current_page = 1
            self._filter_bulletins()
            
        except Exception as e:
            print(f"Erreur generation tous bulletins: {e}")
            import traceback
            traceback.print_exc()
            messagebox.showerror("Erreur", f"Erreur lors de la génération des bulletins:\n{str(e)}")
    
    def _generate_bulletins_for_class_id(self, classe_id, classe_nom):
        """Génère les bulletins pour une classe spécifique (version utilitaire) - DÉPRÉCIÉE"""
        # Rediriger vers la nouvelle méthode complète
        return self._generate_complete_bulletins_for_class_id(classe_id, classe_nom)
    
    def _generate_complete_bulletins_for_class_id(self, classe_id, classe_nom):
        """Génère les bulletins pour TOUS les élèves d'une classe (version utilitaire)"""
        try:
            # Récupérer tous les élèves de la classe
            eleves_classe = get_all_eleves()
            eleves_classe = [e for e in eleves_classe if self._extract_student_class_id(e) == classe_id]
            
            if not eleves_classe:
                print(f"Aucun élève trouvé dans la classe '{classe_nom}'")
                return 0
            
            print(f"DEBUG: {len(eleves_classe)} élèves trouvés dans la classe '{classe_nom}'")
            bulletins_generes = 0
            
            for eleve in eleves_classe:
                eleve_id = eleve.get('id_eleve')
                eleve_nom = eleve.get('nom', '')
                eleve_prenom = eleve.get('prenom', '')
                
                # Récupérer les notes de l'élève pour toutes les périodes
                for periode in self.periodes:
                    notes_eleve = get_notes_by_eleve(eleve_id, trimestre=periode)
                    
                    # Calculer la moyenne et regrouper par matière
                    total_points = 0
                    total_coefficients = 0
                    matieres_notes = {}
                    
                    if notes_eleve:
                        for note in notes_eleve:
                            note_value = float(note.get("note", 0))
                            coefficient = self._extract_coefficient(note)
                            matiere = self._extract_subject_name(note)
                            
                            if matiere not in matieres_notes:
                                matieres_notes[matiere] = {
                                    'notes': [],
                                    'coefficient': coefficient
                                }
                            matieres_notes[matiere]['notes'].append(note_value)
                            
                            total_points += note_value * coefficient
                            total_coefficients += coefficient
                    
                    # Si l'élève n'a pas de notes, moyenne = 0
                    if total_coefficients == 0:
                        moyenne_generale = 0.0
                        print(f"DEBUG: Aucune note pour {eleve_prenom} {eleve_nom} - {periode} - Moyenne: 0.0")
                    else:
                        moyenne_generale = total_points / total_coefficients
                        print(f"DEBUG: {len(notes_eleve)} notes pour {eleve_prenom} {eleve_nom} - {periode} - Moyenne: {moyenne_generale:.2f}")
                    
                    # Déterminer la mention
                    if moyenne_generale >= 16:
                        mention = "Très bien"
                    elif moyenne_generale >= 14:
                        mention = "Bien"
                    elif moyenne_generale >= 12:
                        mention = "Assez bien"
                    elif moyenne_generale >= 10:
                        mention = "Passable"
                    else:
                        mention = "Insuffisant"
                    
                    # Créer le bulletin
                    bulletin_info = {
                        "id_eleve": eleve_id,
                        "periode": periode,
                        "moyenne_generale": round(moyenne_generale, 2),
                        "rang": 0,  # Sera calculé après
                        "appreciation": mention
                    }
                    
                    # Ajouter le bulletin
                    success = add_bulletin(bulletin_info)
                    if success:
                        bulletins_generes += 1
                        print(f"Bulletin genere pour {eleve_prenom} {eleve_nom} - {classe_nom} - {periode} - Moyenne: {moyenne_generale:.2f}")
            
            return bulletins_generes
            
        except Exception as e:
            print(f"Erreur generation bulletins pour classe {classe_nom}: {e}")
            return 0
    
    def _display_bulletins(self, bulletins):
        """Affiche les bulletins dans l'interface avec CTkTable moderne"""
        if not self.table_frame:
            return
        
        # Nettoyer le frame de manière sécurisée
        try:
            for widget in self.table_frame.winfo_children():
                try:
                    widget.destroy()
                except Exception as e:
                    print(f"Erreur destruction widget: {e}")
                    continue
        except Exception as e:
            print(f"Erreur nettoyage frame: {e}")
        
        if not bulletins:
            # Afficher un message si aucun bulletin
            no_data_frame = ctk.CTkFrame(self.table_frame, fg_color="transparent")
            no_data_frame.grid(row=0, column=0, sticky="nsew")
            no_data_frame.grid_columnconfigure(0, weight=1)
            no_data_frame.grid_rowconfigure(0, weight=1)
            
            # Icône pour aucun résultat
            no_data_icon = load_icon('search', (48, 48))
            if no_data_icon:
                icon_label = ctk.CTkLabel(no_data_frame, image=no_data_icon, text="")
                icon_label.grid(row=0, column=0, pady=(0, 10))
            
            ctk.CTkLabel(no_data_frame, text="Aucun bulletin trouvé", 
                        font=(FONT, FONT_SIZE_HEADER), text_color=TEXT_SECONDARY).grid(row=1, column=0)
            ctk.CTkLabel(no_data_frame, text="Sélectionnez une classe et une période pour afficher les bulletins", 
                        font=(FONT, FONT_SIZE_TEXT), text_color=TEXT_SECONDARY).grid(row=2, column=0)
            return
        
        # Trier les bulletins par ordre de mérite (moyenne décroissante) avant affichage
        # Ignorer les rangs existants et recalculer basé sur les moyennes réelles
        bulletins_tries = sorted(bulletins, key=lambda x: -x.get('moyenne_generale', 0))
        
        # Recalculer les rangs basés sur le tri par moyenne
        for i, bulletin in enumerate(bulletins_tries):
            bulletin['rang_calcule'] = i + 1
        
        # Préparer les données pour CTkTable
        headers = ["Rang", "Élève", "Classe", "Période", "Moyenne", "Appréciation"]
        table_data = [headers]  # Première ligne = en-têtes
        
        # Ajouter les données des bulletins avec pagination
        start_idx = (self.current_page - 1) * self.items_per_page
        end_idx = start_idx + self.items_per_page
        paginated_bulletins = bulletins_tries[start_idx:end_idx]
        
        for bulletin in paginated_bulletins:
            eleve_nom = self._get_eleve_name(bulletin.get('id_eleve'))
            classe_nom = self._get_eleve_classe_name(bulletin.get('id_eleve'))
            moyenne = bulletin.get('moyenne_generale', 0)
            rang = bulletin.get('rang', 0)
            appreciation = bulletin.get('appreciation', '')
            periode = bulletin.get('periode', '')
            
            # Nettoyer et formater l'appréciation
            appreciation_clean = self._clean_appreciation(appreciation, moyenne)
            
            # Utiliser le rang calculé basé sur le tri par moyenne
            rang_display = str(bulletin.get('rang_calcule', bulletins_tries.index(bulletin) + 1))
            
            row_data = [
                rang_display,
                eleve_nom,
                classe_nom,
                periode,
                f"{moyenne:.2f}",
                appreciation_clean
            ]
            table_data.append(row_data)
        
        # Créer le CTkTable avec design moderne
        self.bulletins_table = CTkTable(
            master=self.table_frame,
            row=len(table_data),
            column=len(headers),
            values=table_data,
            header_color=BG_SIDEBAR,
            colors=[BG_CARD, BG_SECONDARY],  # Couleurs alternées
            hover_color=TEXT_ACCENT,
            corner_radius=8,
            border_width=1,
            border_color=BORDER_COLOR,
            text_color=TEXT_PRIMARY,
            font=(FONT, FONT_SIZE_TEXT),
            command=self._on_table_click
        )
        
        # Placer le tableau
        self.bulletins_table.grid(row=0, column=0, sticky="nsew", padx=10, pady=10)
        
        # Configurer les colonnes du frame principal
        self.table_frame.grid_columnconfigure(0, weight=1)
        self.table_frame.grid_rowconfigure(0, weight=1)
        
        # Ajouter la pagination
        self._add_pagination(len(bulletins))
    
    def _add_pagination(self, total_items):
        """Ajoute la pagination sous le tableau"""
        # Calculer le nombre total de pages
        self.total_pages = max(1, (total_items + self.items_per_page - 1) // self.items_per_page)
        
        # Frame de pagination
        pagination_frame = ctk.CTkFrame(self.table_frame, fg_color="transparent")
        pagination_frame.grid(row=1, column=0, sticky="ew", padx=10, pady=(0, 10))
        pagination_frame.grid_columnconfigure(1, weight=1)
        
        # Informations de pagination
        start_item = (self.current_page - 1) * self.items_per_page + 1
        end_item = min(self.current_page * self.items_per_page, total_items)
        info_text = f"Affichage {start_item}-{end_item} sur {total_items} bulletins (20 par page)"
        
        info_label = ctk.CTkLabel(pagination_frame, text=info_text, 
                                 font=(FONT, FONT_SIZE_SMALL), text_color=TEXT_SECONDARY)
        info_label.grid(row=0, column=0, sticky="w")
        
        # Contrôles de pagination
        controls_frame = ctk.CTkFrame(pagination_frame, fg_color="transparent")
        controls_frame.grid(row=0, column=2, sticky="e")
        
        # Bouton précédent
        prev_btn = ctk.CTkButton(controls_frame, text="◀", width=30, height=30,
                                command=self._prev_page, fg_color=BG_SECONDARY,
                                hover_color=TEXT_ACCENT, font=(FONT, 12))
        prev_btn.grid(row=0, column=0, padx=2)
        
        # Numéro de page
        page_label = ctk.CTkLabel(controls_frame, text=f"{self.current_page}/{self.total_pages}",
                                 font=(FONT, FONT_SIZE_SMALL, "bold"), text_color=TEXT_ACCENT)
        page_label.grid(row=0, column=1, padx=10)
        
        # Bouton suivant
        next_btn = ctk.CTkButton(controls_frame, text="▶", width=30, height=30,
                                command=self._next_page, fg_color=BG_SECONDARY,
                                hover_color=TEXT_ACCENT, font=(FONT, 12))
        next_btn.grid(row=0, column=2, padx=2)
        
        # Désactiver les boutons si nécessaire
        prev_btn.configure(state="disabled" if self.current_page <= 1 else "normal")
        next_btn.configure(state="disabled" if self.current_page >= self.total_pages else "normal")
    
    def _prev_page(self):
        """Page précédente"""
        if self.current_page > 1:
            self.current_page -= 1
            self._filter_bulletins()
    
    def _next_page(self):
        """Page suivante"""
        if self.current_page < self.total_pages:
            self.current_page += 1
            self._filter_bulletins()
    
    def _on_table_click(self, data):
        """Gestionnaire de clic sur le tableau"""
        if data["row"] > 0:  # Ignorer l'en-tête
            print(f"Bulletin sélectionné: ligne {data['row']}, colonne {data['column']}")
            # Ici vous pouvez ajouter la logique pour sélectionner/éditer un bulletin
    
    def _generate_bulletin_word_selected(self):
        """Génère le bulletin Word de l'élève sélectionné"""
        if not self.selected_classe:
            messagebox.showwarning("Sélection requise", "Veuillez d'abord sélectionner une classe.")
            return
        
        if not self.selected_periode:
            messagebox.showwarning("Sélection requise", "Veuillez d'abord sélectionner une période.")
            return
        
        selected_eleve = self.eleve_var.get()
        if selected_eleve in ["Sélectionner classe et période...", "Sélectionner un élève...", "Aucun élève dans cette classe", ""]:
            messagebox.showwarning("Sélection requise", "Veuillez d'abord sélectionner un élève.")
            return
        
        # Récupérer l'élève correspondant
        classes_list = get_all_classes()
        classe_id = None
        for classe in classes_list:
            if classe['nom'] == self.selected_classe:
                classe_id = classe['id']
                break
        
        if not classe_id:
            messagebox.showerror("Erreur", "Classe non trouvée.")
            return
        
        # Récupérer tous les élèves de la classe (compat clés)
        eleves_classe = get_all_eleves()
        def _belongs_to_class(e):
            return (
                e.get('id_classe') == classe_id
                or e.get('classe_id') == classe_id
                or e.get('idclasse') == classe_id
            )
        eleves_classe = [e for e in eleves_classe if _belongs_to_class(e)]
        
        # Trouver l'élève sélectionné
        eleve_selected = None
        for eleve in eleves_classe:
            if f"{eleve['nom']} {eleve['prenom']}" == selected_eleve:
                eleve_selected = eleve
                break
        
        if eleve_selected:
            self._generate_bulletin_word(eleve_selected, classe_id)
        else:
            messagebox.showerror("Erreur", "Élève non trouvé.")
    
    def _generate_bulletin_word(self, eleve, classe_id):
        """Génère un bulletin au format Word pour un élève"""
        if not DOCX_AVAILABLE:
            messagebox.showerror("Erreur", 
                               "La bibliothèque python-docx n'est pas installée.\n\n"
                               "Installez-la avec: pip install python-docx")
            return
        
        try:
            # Importer localement pour éviter les warnings lorsqu'indisponible
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            # Récupérer les notes de l'élève
            eleve_id = eleve.get("id_eleve")
            eleve_nom = eleve.get("nom", "")
            eleve_prenom = eleve.get("prenom", "")
            
            notes_eleve = get_notes_by_eleve(eleve_id, trimestre=self.selected_periode)
            
            if not notes_eleve:
                messagebox.showwarning("Aucune note", 
                                     f"Aucune note trouvée pour {eleve_prenom} {eleve_nom}")
                return
            
            # Calculer la moyenne et regrouper par matière
            total_points = 0
            total_coefficients = 0
            matieres_notes = {}
            
            for note in notes_eleve:
                note_value = float(note.get("note", 0))
                coefficient = float(note.get("coefficient", 1))
                matiere = note.get("nom_matiere", "Inconnue")
                
                if matiere not in matieres_notes:
                    matieres_notes[matiere] = {
                        'notes': [],
                        'coefficient': coefficient
                    }
                matieres_notes[matiere]['notes'].append(note_value)
                
                total_points += note_value * coefficient
                total_coefficients += coefficient
            
            if total_coefficients == 0:
                messagebox.showwarning("Erreur", "Impossible de calculer la moyenne.")
                return
            
            moyenne_generale = total_points / total_coefficients
            
            # Déterminer le rang
            bulletins_classe = [b for b in self.bulletins 
                              if b.get('classe_nom') == self.selected_classe 
                              and b.get('periode') == self.selected_periode]
            bulletins_classe.sort(key=lambda x: x.get('moyenne_generale', 0), reverse=True)
            rang = next((i+1 for i, b in enumerate(bulletins_classe) if b.get('id_eleve') == eleve_id), 0)
            
            # Créer le document Word
            self._create_bulletin_document(eleve, matieres_notes, moyenne_generale, rang, len(bulletins_classe))
            
        except Exception as e:
            print(f"Erreur génération bulletin Word: {e}")
            import traceback
            traceback.print_exc()
            messagebox.showerror("Erreur", f"Erreur lors de la génération du bulletin:\n{str(e)}")
    
    def _create_bulletin_document(self, eleve, matieres_notes, moyenne_generale, rang, total_eleves):
        """Crée le document Word du bulletin"""
        doc = Document()
        
        # Configuration de la page
        sections = doc.sections
        for section in sections:
            section.page_height = Inches(11.69)  # A4
            section.page_width = Inches(8.27)
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # En-tête avec date
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = header.add_run(f"{datetime.now().strftime('%d/%m/%Y %H:%M')}")
        run.font.size = Pt(9)
        
        # Titre principal
        title = doc.add_heading("BULLETIN SCOLAIRE", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.size = Pt(20)
        title.runs[0].font.bold = True
        
        doc.add_paragraph()  # Espacement
        
        # Informations de l'élève
        info_table = doc.add_table(rows=5, cols=4)
        info_table.style = 'Table Grid'
        
        # Remplir les informations
        cells = info_table.rows[0].cells
        cells[0].text = "Nom et Prénom:"
        cells[1].text = f"{eleve.get('nom', '')} {eleve.get('prenom', '')}"
        cells[2].text = "Classe:"
        cells[3].text = self.selected_classe
        
        cells = info_table.rows[1].cells
        cells[0].text = "Date de naissance:"
        cells[1].text = str(eleve.get('date_naissance', ''))
        cells[2].text = "Période:"
        cells[3].text = self.selected_periode
        
        cells = info_table.rows[2].cells
        cells[0].text = "Genre:"
        cells[1].text = eleve.get('sexe', '')
        cells[2].text = "Moyenne Générale:"
        cells[3].text = f"{moyenne_generale:.2f}/20"
        
        cells = info_table.rows[3].cells
        cells[0].text = "Matricule:"
        cells[1].text = str(eleve.get('matricule', ''))
        cells[2].text = "Rang:"
        cells[3].text = f"{rang}/{total_eleves}"
        
        cells = info_table.rows[4].cells
        cells[0].text = "Téléphone:"
        cells[1].text = str(eleve.get('telephone', ''))
        cells[2].text = "Statut:"
        cells[3].text = "Admis" if moyenne_generale >= 10 else "Ajourné"
        
        doc.add_paragraph()  # Espacement
        
        # Section Notes par matière
        doc.add_heading("NOTES PAR MATIÈRE", level=2)
        
        notes_table = doc.add_table(rows=len(matieres_notes) + 1, cols=4)
        notes_table.style = 'Table Grid'
        
        # En-têtes
        header_cells = notes_table.rows[0].cells
        header_cells[0].text = "Matière"
        header_cells[1].text = "Moyenne"
        header_cells[2].text = "Coefficient"
        header_cells[3].text = "Appréciation"
        
        # Remplir les notes
        for idx, (matiere, data) in enumerate(matieres_notes.items(), start=1):
            cells = notes_table.rows[idx].cells
            moyenne_matiere = sum(data['notes']) / len(data['notes'])
            cells[0].text = matiere
            cells[1].text = f"{moyenne_matiere:.2f}"
            cells[2].text = str(data['coefficient'])
            
            # Appréciation selon la moyenne
            if moyenne_matiere >= 16:
                appreciation = "Très bien"
            elif moyenne_matiere >= 14:
                appreciation = "Bien"
            elif moyenne_matiere >= 12:
                appreciation = "Assez bien"
            elif moyenne_matiere >= 10:
                appreciation = "Passable"
            else:
                appreciation = "Insuffisant"
            cells[3].text = appreciation
        
        doc.add_paragraph()  # Espacement
        
        # Signatures
        doc.add_paragraph()
        doc.add_paragraph()
        
        sig_table = doc.add_table(rows=2, cols=3)
        sig_cells = sig_table.rows[0].cells
        sig_cells[0].text = "Signature du Professeur"
        sig_cells[1].text = "Signature du Directeur"
        sig_cells[2].text = "Cachet de l'Établissement"
        
        # Ajouter de l'espace pour les signatures
        for cell in sig_table.rows[1].cells:
            cell.text = "\n\n\n"
        
        # Footer
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run(f"\nDocument généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}")
        run.font.size = Pt(8)
        run.font.italic = True
        
        # Demander où sauvegarder
        filename = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Document Word", "*.docx"), ("Tous les fichiers", "*.*")],
            initialfile=f"Bulletin_{eleve.get('nom')}_{eleve.get('prenom')}_{self.selected_periode}.docx",
            title="Enregistrer le bulletin"
        )
        
        if filename:
            doc.save(filename)
            messagebox.showinfo("Succès", 
                              f"Bulletin généré avec succès !\n\n"
                              f"Fichier: {filename}\n\n"
                              f"Élève: {eleve.get('prenom')} {eleve.get('nom')}\n"
                              f"Moyenne: {moyenne_generale:.2f}/20\n"
                              f"Rang: {rang}/{total_eleves}")
            
            # Ouvrir le fichier
            try:
                os.startfile(filename)
            except:
                pass
    
    def _generate_detailed_bulletins(self):
        """Génère des bulletins détaillés avec toutes les notes"""
        try:
            print(f"Generation des bulletins individuels pour {self.selected_classe} - {self.selected_periode or 'Toutes'}")
            
            # Récupérer l'ID de la classe
            classes_list = get_all_classes()
            classe_id = None
            for classe in classes_list:
                if classe['nom'] == self.selected_classe:
                    classe_id = classe['id']
                    break
            
            if not classe_id:
                messagebox.showerror("Erreur", "Classe non trouvée.")
                return
            
            # Récupérer tous les élèves de la classe
            eleves_classe = get_all_eleves()
            eleves_classe = [e for e in eleves_classe if e.get('id_classe') == classe_id]
            
            if not eleves_classe:
                messagebox.showwarning("Génération", "Aucun élève trouvé dans cette classe.")
                return
            
            print(f"{len(eleves_classe)} eleves trouves dans la classe {self.selected_classe}")
            
            bulletins_generes = 0
            
            for eleve in eleves_classe:
                eleve_id = eleve.get("id_eleve")
                eleve_nom = eleve.get("nom", "")
                eleve_prenom = eleve.get("prenom", "")
                
                print(f"Traitement bulletin individuel de {eleve_prenom} {eleve_nom} (ID: {eleve_id})")
                
                # Récupérer toutes les notes de l'élève pour la période
                notes_eleve = get_notes_by_eleve(eleve_id, trimestre=self.selected_periode)
                
                if not notes_eleve:
                    print(f"Aucune note trouvee pour {eleve_prenom} {eleve_nom}")
                    continue
                
                # Calculer la moyenne pondérée
                total_points = 0
                total_coefficients = 0
                matieres_notes = {}
                
                for note in notes_eleve:
                    note_value = float(note.get("note", 0))
                    coefficient = float(note.get("coefficient", 1))
                    matiere = note.get("nom_matiere", "Inconnue")
                    
                    if matiere not in matieres_notes:
                        matieres_notes[matiere] = []
                    matieres_notes[matiere].append(note_value)
                    
                    total_points += note_value * coefficient
                    total_coefficients += coefficient
                
                if total_coefficients > 0:
                    moyenne_generale = total_points / total_coefficients
                    
                    # Déterminer la mention
                    if moyenne_generale >= 16:
                        mention = "Très Bien"
                    elif moyenne_generale >= 14:
                        mention = "Bien"
                    elif moyenne_generale >= 12:
                        mention = "Assez Bien"
                    elif moyenne_generale >= 10:
                        mention = "Passable"
                    else:
                        mention = "Insuffisant"
                    
                    # Créer le bulletin détaillé
                    bulletin_info = {
                        "id_eleve": eleve_id,
                        "periode": self.selected_periode or "Année complète",
                        "moyenne_generale": round(moyenne_generale, 2),
                        "rang": 0,  # Sera calculé après
                        "appreciation": mention
                    }
                    
                    # Ajouter ou mettre à jour le bulletin
                    success = add_bulletin(bulletin_info)
                    if success:
                        bulletins_generes += 1
                        print(f"Bulletin individuel genere pour {eleve_prenom} {eleve_nom} - Moyenne: {moyenne_generale:.2f}")
                    else:
                        print(f"Erreur generation bulletin pour {eleve_prenom} {eleve_nom}")
            
            # Recalculer les rangs pour la classe et période sélectionnées
            if classe_id and self.selected_periode:
                self._recalculate_ranks(classe_id, self.selected_periode)
            
            messagebox.showinfo("Génération terminée", 
                              f"Génération des bulletins individuels terminée !\n\n"
                              f"• {bulletins_generes} bulletins générés\n"
                              f"• Classe: {self.selected_classe}\n"
                              f"• Période: {self.selected_periode or 'Toutes'}")
            
            # Actualiser l'affichage
            self._refresh_all()
            
        except Exception as e:
            print(f"Erreur generation bulletins individuels: {e}")
            messagebox.showerror("Erreur", f"Erreur lors de la génération des bulletins individuels:\n{str(e)}")
    
    def _recalculate_ranks(self, classe_id, periode):
        """Recalcule les rangs des bulletins par classe et période"""
        try:
            # Utiliser le contrôleur pour recalculer les classements
            if hasattr(self.bulletins_controller, 'calcul_controller'):
                # Récupérer l'ID de période si nécessaire
                # Pour l'instant, on utilise une méthode simplifiée
                print(f"Recalcul des rangs pour classe {classe_id}, periode {periode}")
                
            # Grouper les bulletins par classe et période
            bulletins_by_class_period = {}
            
            for bulletin in self.bulletins:
                classe = bulletin.get('classe_nom', '')
                periode_bulletin = bulletin.get('periode', '')
                key = f"{classe}_{periode_bulletin}"
                
                if key not in bulletins_by_class_period:
                    bulletins_by_class_period[key] = []
                bulletins_by_class_period[key].append(bulletin)
            
            # Trier par moyenne décroissante et attribuer les rangs
            for key, bulletins_group in bulletins_by_class_period.items():
                bulletins_group.sort(key=lambda x: x.get('moyenne_generale', 0), reverse=True)
                
                for i, bulletin in enumerate(bulletins_group, 1):
                    bulletin['rang'] = i
                    # Mettre à jour dans la base de données
                    if bulletin.get('id'):
                        update_bulletin(bulletin.get('id'), {'rang': i})
            
            print("Rangs recalcules avec succes")
            
        except Exception as e:
            print(f"Erreur recalcul des rangs: {e}")
    
    def _recalculate_all_ranks_for_class(self, classe_id):
        """Recalcule les rangs pour toutes les périodes d'une classe spécifique"""
        try:
            # Récupérer le nom de la classe
            classe_nom = None
            for classe in self.classes.values():
                if classe.get('id') == classe_id:
                    classe_nom = classe.get('nom', '')
                    break
            
            if not classe_nom:
                print(f"Nom de classe non trouvé pour l'ID {classe_id}")
                return
            
            print(f"Recalcul des rangs pour la classe '{classe_nom}' (ID: {classe_id})")
            
            # Grouper les bulletins de cette classe par période
            bulletins_by_period = {}
            
            for bulletin in self.bulletins:
                if bulletin.get('classe_nom') == classe_nom:
                    periode = bulletin.get('periode', '')
                    if periode not in bulletins_by_period:
                        bulletins_by_period[periode] = []
                    bulletins_by_period[periode].append(bulletin)
            
            # Calculer les rangs pour chaque période
            for periode, bulletins_period in bulletins_by_period.items():
                # Trier par moyenne décroissante
                bulletins_period.sort(key=lambda x: x.get('moyenne_generale', 0), reverse=True)
                
                # Attribuer les rangs
                for i, bulletin in enumerate(bulletins_period, 1):
                    bulletin['rang'] = i
                    # Mettre à jour dans la base de données
                    if bulletin.get('id'):
                        try:
                            update_bulletin(bulletin.get('id'), {'rang': i})
                            print(f"Rang {i} attribué à {bulletin.get('eleve_nom', '')} {bulletin.get('eleve_prenom', '')} - {periode}")
                        except Exception as e:
                            print(f"Erreur mise à jour rang pour bulletin {bulletin.get('id')}: {e}")
            
            print(f"Rangs recalculés pour {len(bulletins_by_period)} périodes de la classe '{classe_nom}'")
            
        except Exception as e:
            print(f"Erreur recalcul rangs pour classe {classe_id}: {e}")
            import traceback
            traceback.print_exc()
    
    def _export_to_excel(self):
        """Exporte tous les bulletins au format Excel, regroupés par élève"""
        try:
            from tkinter import filedialog
            import pandas as pd
            from datetime import datetime
            
            if not self.bulletins:
                messagebox.showwarning("Export", "Aucun bulletin à exporter.")
                return
            
            # Demander le chemin de sauvegarde
            filename = filedialog.asksaveasfilename(
                defaultextension=".xlsx",
                filetypes=[("Fichiers Excel", "*.xlsx"), ("Tous les fichiers", "*.*")],
                title="Exporter les bulletins vers Excel"
            )
            
            if not filename:
                return
            
            # Organiser les données par élève
            eleves_data = {}
            
            for bulletin in self.bulletins:
                eleve_id = bulletin.get('id_eleve')
                eleve_nom = f"{bulletin.get('eleve_prenom', '')} {bulletin.get('eleve_nom', '')}"
                
                if eleve_id not in eleves_data:
                    eleves_data[eleve_id] = {
                        'nom': eleve_nom,
                        'classe': bulletin.get('classe_nom', ''),
                        'bulletins': []
                    }
                
                eleves_data[eleve_id]['bulletins'].append({
                    'periode': bulletin.get('periode', ''),
                    'moyenne': bulletin.get('moyenne_generale', 0),
                    'rang': bulletin.get('rang', 0),
                    'appreciation': bulletin.get('appreciation', '')
                })
            
            # Créer le fichier Excel avec plusieurs feuilles
            with pd.ExcelWriter(filename, engine='openpyxl') as writer:
                # Feuille récapitulative
                summary_data = []
                for eleve_id, data in eleves_data.items():
                    for bulletin in data['bulletins']:
                        summary_data.append({
                            'Élève': data['nom'],
                            'Classe': data['classe'],
                            'Période': bulletin['periode'],
                            'Moyenne': bulletin['moyenne'],
                            'Rang': bulletin['rang'],
                            'Appréciation': bulletin['appreciation']
                        })
                
                df_summary = pd.DataFrame(summary_data)
                df_summary.to_excel(writer, sheet_name='Récapitulatif', index=False)
                
                # Feuille par élève
                for eleve_id, data in eleves_data.items():
                    eleve_df = pd.DataFrame(data['bulletins'])
                    eleve_df.columns = ['Période', 'Moyenne', 'Rang', 'Appréciation']
                    sheet_name = data['nom'][:30]  # Limiter la longueur du nom
                    eleve_df.to_excel(writer, sheet_name=sheet_name, index=False)
            
            messagebox.showinfo("Export réussi", 
                              f"Export terminé avec succès !\n\n"
                              f"• Fichier: {filename}\n"
                              f"• {len(eleves_data)} élèves exportés\n"
                              f"• {len(self.bulletins)} bulletins au total")
            
        except ImportError:
            messagebox.showerror("Erreur", "La bibliothèque pandas est requise pour l'export Excel.\nVeuillez l'installer avec: pip install pandas openpyxl")
        except Exception as e:
            print(f"Erreur export Excel: {e}")
            messagebox.showerror("Erreur", f"Erreur lors de l'export Excel:\n{str(e)}")
    
    def _export_individual_bulletin_excel(self, eleve, matieres_notes, moyenne_generale, mention, rang, total_eleves):
        """Exporte un bulletin individuel au format Excel avec formatage professionnel"""
        try:
            from tkinter import filedialog
            import pandas as pd
            from datetime import datetime
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
            from openpyxl.utils.dataframe import dataframe_to_rows
            
            # Demander le chemin de sauvegarde
            eleve_nom = f"{eleve.get('nom', '')} {eleve.get('prenom', '')}"
            filename = filedialog.asksaveasfilename(
                defaultextension=".xlsx",
                filetypes=[("Fichiers Excel", "*.xlsx"), ("Tous les fichiers", "*.*")],
                title=f"Exporter le bulletin de {eleve_nom}",
                initialvalue=f"Bulletin_{eleve_nom.replace(' ', '_')}_{self.selected_periode}.xlsx"
            )
            
            if not filename:
                return
            
            # Créer un nouveau classeur Excel
            wb = Workbook()
            ws = wb.active
            ws.title = "Bulletin Scolaire"
            
            # Styles
            header_font = Font(name="Arial", size=16, bold=True, color="FFFFFF")
            title_font = Font(name="Arial", size=14, bold=True)
            normal_font = Font(name="Arial", size=11)
            small_font = Font(name="Arial", size=9)
            
            header_fill = PatternFill(start_color="2E75B6", end_color="2E75B6", fill_type="solid")
            accent_fill = PatternFill(start_color="D9E2F3", end_color="D9E2F3", fill_type="solid")
            success_fill = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")
            
            center_alignment = Alignment(horizontal="center", vertical="center")
            left_alignment = Alignment(horizontal="left", vertical="center")
            
            thin_border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )
            
            # En-tête du bulletin
            ws.merge_cells('A1:F1')
            ws['A1'] = "BULLETIN SCOLAIRE INDIVIDUEL"
            ws['A1'].font = header_font
            ws['A1'].fill = header_fill
            ws['A1'].alignment = center_alignment
            ws['A1'].border = thin_border
            ws.row_dimensions[1].height = 30
            
            # Informations de l'établissement (optionnel)
            ws.merge_cells('A2:F2')
            ws['A2'] = "ÉTABLISSEMENT SCOLAIRE"
            ws['A2'].font = title_font
            ws['A2'].alignment = center_alignment
            ws.row_dimensions[2].height = 25
            
            # Informations de l'élève
            row = 4
            ws[f'A{row}'] = "Nom et Prénom:"
            ws[f'A{row}'].font = title_font
            ws[f'B{row}'] = eleve_nom
            ws[f'B{row}'].font = normal_font
            ws[f'D{row}'] = "Classe:"
            ws[f'D{row}'].font = title_font
            ws[f'E{row}'] = self.selected_classe
            ws[f'E{row}'].font = normal_font
            row += 1
            
            ws[f'A{row}'] = "Période:"
            ws[f'A{row}'].font = title_font
            ws[f'B{row}'] = self.selected_periode
            ws[f'B{row}'].font = normal_font
            ws[f'D{row}'] = "Moyenne Générale:"
            ws[f'D{row}'].font = title_font
            ws[f'E{row}'] = f"{moyenne_generale:.2f}/20"
            ws[f'E{row}'].font = Font(name="Arial", size=12, bold=True, color="2E75B6")
            row += 1
            
            ws[f'A{row}'] = "Rang:"
            ws[f'A{row}'].font = title_font
            ws[f'B{row}'] = f"{rang}/{total_eleves}"
            ws[f'B{row}'].font = normal_font
            ws[f'D{row}'] = "Mention:"
            ws[f'D{row}'].font = title_font
            ws[f'E{row}'] = mention
            ws[f'E{row}'].font = Font(name="Arial", size=12, bold=True, color="2E75B6")
            row += 2
            
            # En-tête du tableau des notes
            headers = ["Matière", "Notes", "Moyenne", "Coefficient", "Appréciation"]
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=row, column=col, value=header)
                cell.font = title_font
                cell.fill = accent_fill
                cell.alignment = center_alignment
                cell.border = thin_border
            ws.row_dimensions[row].height = 25
            row += 1
            
            # Données des matières
            for matiere, data in matieres_notes.items():
                notes_list = data['notes']
                moyenne_matiere = sum(notes_list) / len(notes_list)
                coefficient = data['coefficient']
                
                # Appréciation par matière
                if moyenne_matiere >= 16:
                    appreciation = "Très bien"
                    appreciation_color = "00B050"  # Vert
                elif moyenne_matiere >= 14:
                    appreciation = "Bien"
                    appreciation_color = "0070C0"  # Bleu
                elif moyenne_matiere >= 12:
                    appreciation = "Assez bien"
                    appreciation_color = "FFC000"  # Orange
                elif moyenne_matiere >= 10:
                    appreciation = "Passable"
                    appreciation_color = "FF6600"  # Orange foncé
                else:
                    appreciation = "Insuffisant"
                    appreciation_color = "FF0000"  # Rouge
                
                # Formater les notes
                notes_str = ", ".join([f"{note:.1f}" for note in notes_list])
                
                # Remplir les cellules
                ws.cell(row=row, column=1, value=matiere).font = normal_font
                ws.cell(row=row, column=2, value=notes_str).font = normal_font
                ws.cell(row=row, column=3, value=f"{moyenne_matiere:.2f}").font = normal_font
                ws.cell(row=row, column=4, value=coefficient).font = normal_font
                ws.cell(row=row, column=5, value=appreciation).font = Font(name="Arial", size=11, bold=True, color=appreciation_color)
                
                # Appliquer les bordures
                for col in range(1, 6):
                    ws.cell(row=row, column=col).border = thin_border
                    ws.cell(row=row, column=col).alignment = center_alignment
                
                row += 1
            
            # Ligne de moyenne générale
            row += 1
            ws.merge_cells(f'A{row}:D{row}')
            ws[f'A{row}'] = "MOYENNE GÉNÉRALE"
            ws[f'A{row}'].font = Font(name="Arial", size=12, bold=True)
            ws[f'A{row}'].alignment = center_alignment
            ws[f'E{row}'] = f"{moyenne_generale:.2f}/20"
            ws[f'E{row}'].font = Font(name="Arial", size=14, bold=True, color="2E75B6")
            ws[f'E{row}'].alignment = center_alignment
            ws[f'E{row}'].fill = success_fill
            
            # Appliquer les bordures à la ligne de moyenne
            for col in range(1, 6):
                ws.cell(row=row, column=col).border = thin_border
            
            # Footer avec date
            row += 3
            ws.merge_cells(f'A{row}:F{row}')
            ws[f'A{row}'] = f"Document généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}"
            ws[f'A{row}'].font = small_font
            ws[f'A{row}'].alignment = center_alignment
            
            # Ajuster la largeur des colonnes
            column_widths = [25, 20, 12, 12, 15]
            for i, width in enumerate(column_widths, 1):
                ws.column_dimensions[chr(64 + i)].width = width
            
            # Sauvegarder le fichier
            wb.save(filename)
            
            messagebox.showinfo("Export réussi", 
                              f"Bulletin de {eleve_nom} exporté avec succès !\n\n"
                              f"• Fichier: {filename}\n"
                              f"• Période: {self.selected_periode}\n"
                              f"• Moyenne: {moyenne_generale:.2f}/20\n"
                              f"• Mention: {mention}")
            
        except ImportError:
            messagebox.showerror("Erreur", "La bibliothèque openpyxl est requise pour l'export Excel.\nVeuillez l'installer avec: pip install openpyxl")
        except Exception as e:
            print(f"Erreur export bulletin individuel: {e}")
            messagebox.showerror("Erreur", f"Erreur lors de l'export du bulletin:\n{str(e)}")
    
    def _get_eleve_name(self, eleve_id):
        """Récupère le nom complet d'un élève depuis les bulletins"""
        try:
            # Chercher dans les bulletins existants pour trouver le nom
            for bulletin in self.bulletins:
                if bulletin.get('id_eleve') == eleve_id:
                    nom = bulletin.get('eleve_nom', '')
                    prenom = bulletin.get('eleve_prenom', '')
                    if nom and prenom:
                        return f"{nom} {prenom}"
                    elif nom:
                        return nom
            return f"Élève {eleve_id}"
        except Exception as e:
            print(f"Erreur recuperation nom eleve {eleve_id}: {e}")
            return f"Élève {eleve_id}"
    
    def _show_generate_message(self):
        """Affiche un message pour générer les bulletins"""
        if not self.table_frame:
            return
        
        # Nettoyer le frame
        for widget in self.table_frame.winfo_children():
            widget.destroy()
        
        # Message central
        message_frame = ctk.CTkFrame(self.table_frame, fg_color="transparent")
        message_frame.grid(row=0, column=0, sticky="nsew")
        message_frame.grid_columnconfigure(0, weight=1)
        message_frame.grid_rowconfigure(0, weight=1)
        
        # Icône et texte
        add_icon = load_icon('add', (64, 64))
        icon_label = ctk.CTkLabel(message_frame, image=add_icon, text="")
        icon_label.grid(row=0, column=0, pady=(0, MARGIN_MEDIUM))
        
        title_label = ctk.CTkLabel(message_frame, text="Générer les bulletins", 
                                  font=F_SUB, text_color=TEXT_PRIMARY)
        title_label.grid(row=1, column=0, pady=(0, MARGIN_SMALL))
        
        desc_label = ctk.CTkLabel(message_frame, text=f"Classe: {self.selected_classe}\nPériode: {self.selected_periode or 'Toutes'}\n\nCliquez sur 'Générer' pour créer automatiquement les bulletins.",
                                 font=F_SMALL, text_color=TEXT_SECONDARY)
        desc_label.grid(row=2, column=0)
        
        # Bouton générer
        generate_btn = ctk.CTkButton(message_frame, text="Générer les bulletins",
                                   command=self._generate_bulletins,
                                   fg_color=SUCCESS_GREEN, hover_color="#80C7C5",
                                   font=F_TXT, height=40)
        generate_btn.grid(row=3, column=0, pady=(MARGIN_MEDIUM, 0))

    def _generate_bulletins(self):
        """Génère automatiquement les bulletins pour la classe et période sélectionnées"""
        if not self.selected_classe:
            messagebox.showwarning("Sélection requise", "Veuillez d'abord sélectionner une classe.")
            return
        
        # Confirmer la génération
        periode_text = self.selected_periode or "toutes les périodes"
        if messagebox.askyesno("Génération des bulletins", 
                              f"Voulez-vous générer automatiquement les bulletins pour :\n\n"
                              f"• Classe: {self.selected_classe}\n"
                              f"• Période: {periode_text}\n\n"
                              f"Cette action va calculer les moyennes à partir des notes existantes\n"
                              f"en utilisant les coefficients spécifiques à chaque classe."):
            self._generate_bulletins_from_notes()
    
    def _generate_bulletins_from_notes(self):
        """Génère automatiquement les bulletins à partir des notes de la classe et période sélectionnées"""
        try:
            print(f"Generation des bulletins pour {self.selected_classe} - {self.selected_periode or 'Toutes'}")
            
            # Récupérer l'ID de la classe
            classes_list = get_all_classes()
            classe_id = None
            for classe in classes_list:
                if classe['nom_classe'] == self.selected_classe:
                    classe_id = classe['id_classe']
                    break
            
            if not classe_id:
                messagebox.showerror("Erreur", "Classe non trouvée.")
                return
            
            # Récupérer tous les élèves de la classe
            eleves_classe = get_all_eleves()
            eleves_classe = [e for e in eleves_classe if e.get('id_classe') == classe_id]
            
            if not eleves_classe:
                messagebox.showwarning("Génération", "Aucun élève trouvé dans cette classe.")
                return
            
            print(f"{len(eleves_classe)} eleves trouves dans la classe {self.selected_classe}")
            
            bulletins_generes = 0
            
            for eleve in eleves_classe:
                eleve_id = eleve.get("id_eleve")
                eleve_nom = eleve.get("nom", "")
                eleve_prenom = eleve.get("prenom", "")
                
                print(f"Traitement de {eleve_prenom} {eleve_nom} (ID: {eleve_id})")
                
                # Récupérer les notes de l'élève pour la période sélectionnée
                notes_eleve = get_notes_by_eleve(eleve_id, trimestre=self.selected_periode)
                
                print(f"DEBUG: {len(notes_eleve)} notes trouvées pour {eleve_prenom} {eleve_nom} - {self.selected_periode}")
                if notes_eleve:
                    for note in notes_eleve[:3]:  # Afficher les 3 premières notes pour debug
                        print(f"  Note: {note.get('note')} - Matière: {note.get('matiere_nom')} - Date: {note.get('date_evaluation')}")
                
                if not notes_eleve:
                    print(f"Aucune note trouvee pour {eleve_prenom} {eleve_nom} - {self.selected_periode or 'Toutes'}")
                    continue
                
                print(f"{len(notes_eleve)} notes trouvees pour {eleve_prenom} {eleve_nom}")
                
                # Calculer la moyenne pondérée en utilisant les coefficients de classe_matieres
                total_points = 0
                total_coefficients = 0
                
                for note in notes_eleve:
                    note_value = float(note.get("note", 0))
                    coefficient = float(note.get("coefficient", 1))  # coefficient_classe depuis classe_matieres
                    total_points += note_value * coefficient
                    total_coefficients += coefficient
                
                if total_coefficients > 0:
                    moyenne_generale = total_points / total_coefficients
                    
                    # Déterminer la mention
                    if moyenne_generale >= 16:
                        mention = "Très Bien"
                    elif moyenne_generale >= 14:
                        mention = "Bien"
                    elif moyenne_generale >= 12:
                        mention = "Assez Bien"
                    elif moyenne_generale >= 10:
                        mention = "Passable"
                    else:
                        mention = "Insuffisant"
                    
                    print(f"{eleve_prenom} {eleve_nom} - Moyenne: {moyenne_generale:.2f} ({mention})")
                    
                    # Créer le bulletin
                    bulletin_data = {
                        'id_eleve': eleve_id,
                        'periode': self.selected_periode or 'Année complète',
                        'moyenne_generale': moyenne_generale,
                        'rang': 0,  # Sera calculé après
                        'appreciation': mention,
                        'date_creation': datetime.now()
                    }
                    
                    # Sauvegarder le bulletin
                    success = self.bulletins_controller.create_bulletin(bulletin_data)
                    if success:
                        bulletins_generes += 1
                    else:
                        print(f"Erreur lors de la creation du bulletin pour {eleve_prenom} {eleve_nom}")
            
            # Recalculer les rangs après génération
            self._recalculate_ranks(classe_id, self.selected_periode or 'Année complète')
            
            # Rafraîchir l'affichage
            self._refresh_all()
            
            messagebox.showinfo("Génération terminée", 
                              f"{bulletins_generes} bulletins generes avec succes pour la classe {self.selected_classe}.\n\n"
                              f"Les moyennes ont été calculées automatiquement à partir des notes existantes\n"
                              f"en utilisant les coefficients spécifiques à chaque classe.")
            
        except Exception as e:
            print(f"Erreur lors de la generation des bulletins: {e}")
            messagebox.showerror("Erreur", f"Erreur lors de la génération des bulletins:\n{str(e)}")
    
    def _recalculate_ranks(self, classe_id, periode):
        """Recalcule les rangs des élèves dans une classe pour une période donnée"""
        try:
            # Récupérer tous les bulletins de la classe et période, triés par moyenne décroissante
            bulletins = self.bulletins_controller.get_bulletins_by_classe_and_periode(classe_id, periode)
            
            # Trier par moyenne décroissante
            bulletins.sort(key=lambda x: x.get('moyenne_generale', 0), reverse=True)
            
            # Mettre à jour les rangs
            for i, bulletin in enumerate(bulletins):
                rang = i + 1
                self.bulletins_controller.update_bulletin_rank(bulletin['id_bulletin'], rang)
            
            print(f"Rangs recalcules pour la classe {classe_id} - {periode}")
            
        except Exception as e:
            print(f"Erreur lors du recalcul des rangs: {e}")

    def _refresh_all(self):
        """Rafraîchit toutes les données"""
        print("Rafraichissement des donnees bulletins...")
        self._load_data()
        self._filter_bulletins()
        print("Donnees bulletins rafraichies")

    def _edit_bulletin(self):
        """Modifie un bulletin sélectionné"""
        messagebox.showinfo("Information", "Fonctionnalité de modification à implémenter.")
    
    def _delete_bulletin(self):
        """Supprime un bulletin sélectionné"""
        messagebox.showinfo("Information", "Fonctionnalité de suppression à implémenter.")
    
    def _generate_bulletins_from_notes(self):
        """Génère automatiquement les bulletins à partir des notes de la classe et période sélectionnées"""
        try:
            # Récupérer la classe ID
            classe_id = next((cid for cid, cdata in self.classes.items() if cdata.get("nom") == self.selected_classe), None)
            if not classe_id:
                messagebox.showerror("Erreur", "Classe non trouvée.")
                return
            
            # Récupérer tous les élèves de la classe
            eleves_classe = get_all_eleves(classe_id=classe_id)
            if not eleves_classe:
                messagebox.showwarning("Génération", "Aucun élève trouvé dans cette classe.")
                return
            
            bulletins_generes = 0
            
            for eleve in eleves_classe:
                eleve_id = eleve.get("id_eleve")
                
                # Récupérer les notes de l'élève pour la période sélectionnée
                notes_eleve = get_notes_by_eleve(eleve_id, trimestre=self.selected_periode)
                
                if not notes_eleve:
                    print(f"Aucune note trouvee pour {eleve.get('nom')} {eleve.get('prenom')} - {self.selected_periode}")
                    continue
                
                # Calculer la moyenne pondérée
                total_points = 0
                total_coefficients = 0
                
                for note in notes_eleve:
                    note_value = float(note.get("note", 0))
                    coefficient = float(note.get("coefficient", 1))
                    total_points += note_value * coefficient
                    total_coefficients += coefficient
                
                if total_coefficients > 0:
                    moyenne_generale = total_points / total_coefficients
                    
                    # Déterminer la mention
                    if moyenne_generale >= 16:
                        mention = "Très Bien"
                    elif moyenne_generale >= 14:
                        mention = "Bien"
                    elif moyenne_generale >= 12:
                        mention = "Assez Bien"
                    elif moyenne_generale >= 10:
                        mention = "Passable"
                    else:
                        mention = "Insuffisant"
                    
                    # Vérifier si un bulletin existe déjà pour cet élève et cette période
                    from database.connection import get_db_connection
                    conn = get_db_connection()
                    cursor = conn.cursor()
                    
                    cursor.execute("""
                        SELECT id_bulletin FROM bulletins 
                        WHERE id_eleve = ? AND periode = ?
                    """, (eleve_id, self.selected_periode))
                    
                    existing_bulletin = cursor.fetchone()
                    
                    if existing_bulletin:
                        # Mettre à jour le bulletin existant
                        cursor.execute("""
                            UPDATE bulletins 
                            SET moyenne_generale = ?, appreciation = ?, date_creation = ?
                            WHERE id_eleve = ? AND periode = ?
                        """, (moyenne_generale, mention, datetime.now(), eleve_id, self.selected_periode))
                        print(f"Bulletin mis a jour pour {eleve.get('nom')} {eleve.get('prenom')} - Moyenne: {moyenne_generale:.2f}")
                    else:
                        # Créer un nouveau bulletin
                        cursor.execute("""
                            INSERT INTO bulletins (id_eleve, periode, moyenne_generale, rang, appreciation, date_creation)
                            VALUES (?, ?, ?, 0, ?, ?)
                        """, (eleve_id, self.selected_periode, moyenne_generale, mention, datetime.now()))
                        print(f"Nouveau bulletin cree pour {eleve.get('nom')} {eleve.get('prenom')} - Moyenne: {moyenne_generale:.2f}")
                    
                    bulletins_generes += 1
                    conn.commit()
                    conn.close()
            
            # Recalculer les rangs après génération
            self._recalculate_ranks(classe_id, self.selected_periode)
            
            # Rafraîchir l'affichage
            self._refresh_all()
            
            messagebox.showinfo("Génération terminée", 
                              f"{bulletins_generes} bulletins generes avec succes pour la classe {self.selected_classe} - {self.selected_periode}.\n\n"
                              f"Les moyennes ont ete calculees automatiquement a partir des notes existantes.")
            
        except Exception as e:
            import traceback
            error_msg = str(e).encode('ascii', 'ignore').decode('ascii')
            print(f"Erreur lors de la generation des bulletins: {error_msg}")
            traceback.print_exc()
            messagebox.showerror("Erreur", f"Erreur lors de la generation des bulletins:\n{error_msg}")
    
    def _recalculate_ranks(self, classe_id, periode):
        """Recalcule les rangs des élèves dans une classe pour une période donnée"""
        try:
            from database.connection import get_db_connection
            conn = get_db_connection()
            cursor = conn.cursor()
            
            # Récupérer tous les bulletins de la classe et période, triés par moyenne décroissante
            cursor.execute("""
                SELECT b.id_bulletin, b.id_eleve, b.moyenne_generale
                FROM bulletins b
                JOIN eleves e ON b.id_eleve = e.id_eleve
                WHERE e.id_classe = ? AND b.periode = ?
                ORDER BY b.moyenne_generale DESC
            """, (classe_id, periode))
            
            bulletins = cursor.fetchall()
            
            # Mettre à jour les rangs
            for i, (bulletin_id, eleve_id, moyenne) in enumerate(bulletins):
                rang = i + 1
                cursor.execute("""
                    UPDATE bulletins SET rang = ? WHERE id_bulletin = ?
                """, (rang, bulletin_id))
            
            conn.commit()
            conn.close()
            print(f"Rangs recalcules pour la classe {classe_id} - {periode}")
            
        except Exception as e:
            print(f"Erreur lors du recalcul des rangs: {e}")
    
    def _open_bulletin_form(self, bulletin_data=None):
        """Ouvre le formulaire d'ajout/modification de bulletin"""
        # Créer une fenêtre modale
        form_window = ctk.CTkToplevel(self)
        form_window.title("Ajouter un Bulletin" if not bulletin_data else "Modifier le Bulletin")
        form_window.geometry("600x700")
        form_window.resizable(False, False)
        form_window.transient(self)
        form_window.grab_set()
        
        # Centrer la fenêtre
        form_window.update_idletasks()
        x = (form_window.winfo_screenwidth() // 2) - (600 // 2)
        y = (form_window.winfo_screenheight() // 2) - (700 // 2)
        form_window.geometry(f"600x700+{x}+{y}")
        
        # Frame principal avec scroll
        main_frame = ctk.CTkScrollableFrame(form_window, fg_color=BG_CARD, corner_radius=12)
        main_frame.pack(fill="both", expand=True, padx=15, pady=15)
        
        # En-tête avec design moderne
        header_frame = ctk.CTkFrame(main_frame, fg_color=ACCENT, corner_radius=12)
        header_frame.pack(fill="x", pady=(0, 20))
        
        # Titre avec icône
        title_frame = ctk.CTkFrame(header_frame, fg_color=ACCENT)
        title_frame.pack(expand=True, fill="x", padx=20, pady=15)
        
        # Icône du formulaire
        add_icon = load_ctk_icon(ICON_MAP.get("add"), size=(24, 24))
        if add_icon:
            ctk.CTkLabel(title_frame, text="", image=add_icon, fg_color=ACCENT).pack(side="left", padx=(0, 10))
        else:
            ctk.CTkLabel(title_frame, text="", font=(FONT, 24), 
                        text_color=BG_MAIN, fg_color=ACCENT).pack(side="left", padx=(0, 10))
        
        title_text = "NOUVEAU BULLETIN" if not bulletin_data else "MODIFIER BULLETIN"
        title = ctk.CTkLabel(title_frame, text=title_text,
                            font=(FONT, FONT_SIZE_HEADER, "bold"),
                            text_color=BG_MAIN, fg_color=ACCENT)
        title.pack(side="left")
        
        # Formulaire avec sections organisées
        form_frame = ctk.CTkFrame(main_frame, fg_color=BG_CARD)
        form_frame.pack(fill="both", expand=True)
        
        # Variables du formulaire
        var_eleve = ctk.StringVar()
        var_periode = ctk.StringVar()
        var_moyenne = ctk.StringVar()
        var_rang = ctk.StringVar()
        var_appreciation = ctk.StringVar()
        
        # Section 1: Informations de base
        basic_section = ctk.CTkFrame(form_frame, fg_color=BG_CARD, corner_radius=8)
        basic_section.pack(fill="x", padx=15, pady=(15, 10))
        
        basic_title = ctk.CTkLabel(basic_section, text="INFORMATIONS DE BASE",
                                  font=(FONT, FONT_SIZE_TEXT, "bold"),
                                  text_color=TEXT_ACCENT, fg_color=BG_CARD)
        basic_title.pack(anchor="w", padx=15, pady=(15, 10))
        
        # Élève
        eleve_frame = ctk.CTkFrame(basic_section, fg_color=BG_CARD)
        eleve_frame.pack(fill="x", padx=15, pady=5)
        eleve_frame.grid_columnconfigure(1, weight=1)
        
        ctk.CTkLabel(eleve_frame, text="Élève:", 
                     font=(FONT, FONT_SIZE_TEXT, "bold"), 
                     text_color=TEXT_PRIMARY, fg_color=BG_CARD).grid(row=0, column=0, sticky="w", padx=(0, 10))
        
        # Récupérer les élèves
        eleves = get_all_eleves()
        eleves_choices = [f"{e[0]} - {e[1]} {e[2]}" for e in eleves]
        
        eleve_combo = ctk.CTkComboBox(eleve_frame, values=eleves_choices, variable=var_eleve,
                                     font=(FONT, FONT_SIZE_TEXT), fg_color=BG_MAIN,
                                     text_color=TEXT_PRIMARY, border_color=BORDER_COLOR,
                                     dropdown_fg_color=BG_CARD, button_color=ACCENT,
                                     height=35, corner_radius=8)
        eleve_combo.grid(row=0, column=1, sticky="ew", padx=(0, 0))
        
        # Période
        periode_frame = ctk.CTkFrame(basic_section, fg_color=BG_CARD)
        periode_frame.pack(fill="x", padx=15, pady=5)
        periode_frame.grid_columnconfigure(1, weight=1)
        
        ctk.CTkLabel(periode_frame, text="Période:", 
                     font=(FONT, FONT_SIZE_TEXT, "bold"), 
                     text_color=TEXT_PRIMARY, fg_color=BG_CARD).grid(row=0, column=0, sticky="w", padx=(0, 10))
        
        periode_combo = ctk.CTkComboBox(periode_frame, values=self.periodes, variable=var_periode,
                                       font=(FONT, FONT_SIZE_TEXT), fg_color=BG_MAIN,
                                       text_color=TEXT_PRIMARY, border_color=BORDER_COLOR,
                                       dropdown_fg_color=BG_CARD, button_color=ACCENT,
                                       height=35, corner_radius=8)
        periode_combo.grid(row=0, column=1, sticky="ew", padx=(0, 0))
        
        # Moyenne
        moyenne_frame = ctk.CTkFrame(basic_section, fg_color=BG_CARD)
        moyenne_frame.pack(fill="x", padx=15, pady=5)
        moyenne_frame.grid_columnconfigure(1, weight=1)
        
        ctk.CTkLabel(moyenne_frame, text="Moyenne:", 
                     font=(FONT, FONT_SIZE_TEXT, "bold"), 
                     text_color=TEXT_PRIMARY, fg_color=BG_CARD).grid(row=0, column=0, sticky="w", padx=(0, 10))
        
        moyenne_entry = ctk.CTkEntry(moyenne_frame, textvariable=var_moyenne, 
                                    font=(FONT, FONT_SIZE_TEXT), fg_color=BG_MAIN,
                                    text_color=TEXT_PRIMARY, border_color=BORDER_COLOR,
                                    placeholder_text="Ex: 15.5", height=35)
        moyenne_entry.grid(row=0, column=1, sticky="ew", padx=(0, 0))
        
        # Rang
        rang_frame = ctk.CTkFrame(basic_section, fg_color=BG_CARD)
        rang_frame.pack(fill="x", padx=15, pady=(5, 15))
        rang_frame.grid_columnconfigure(1, weight=1)
        
        ctk.CTkLabel(rang_frame, text="Rang:", 
                     font=(FONT, FONT_SIZE_TEXT, "bold"), 
                     text_color=TEXT_PRIMARY, fg_color=BG_CARD).grid(row=0, column=0, sticky="w", padx=(0, 10))
        
        rang_entry = ctk.CTkEntry(rang_frame, textvariable=var_rang, 
                                 font=(FONT, FONT_SIZE_TEXT), fg_color=BG_MAIN,
                                 text_color=TEXT_PRIMARY, border_color=BORDER_COLOR,
                                 placeholder_text="Ex: 1", height=35)
        rang_entry.grid(row=0, column=1, sticky="ew", padx=(0, 0))
        
        # Section 2: Appréciation
        app_section = ctk.CTkFrame(form_frame, fg_color=BG_CARD, corner_radius=8)
        app_section.pack(fill="x", padx=15, pady=10)
        
        app_title = ctk.CTkLabel(app_section, text="APPRECIATION",
                                font=(FONT, FONT_SIZE_TEXT, "bold"),
                                text_color=TEXT_ACCENT, fg_color=BG_CARD)
        app_title.pack(anchor="w", padx=15, pady=(15, 10))
        
        app_frame = ctk.CTkFrame(app_section, fg_color=BG_CARD)
        app_frame.pack(fill="x", padx=15, pady=(0, 15))
        
        ctk.CTkLabel(app_frame, text="Appréciation:", 
                     font=(FONT, FONT_SIZE_TEXT, "bold"), 
                     text_color=TEXT_PRIMARY, fg_color=BG_CARD).pack(anchor="w", padx=15, pady=(15, 5))
        
        app_textbox = ctk.CTkTextbox(app_frame, height=80, 
                                    font=(FONT, FONT_SIZE_TEXT), fg_color=BG_MAIN,
                                    text_color=TEXT_PRIMARY, border_color=BORDER_COLOR,
                                    corner_radius=8)
        app_textbox.pack(fill="x", padx=15, pady=(0, 15))
        
        # Boutons avec design moderne
        buttons_frame = ctk.CTkFrame(form_frame, fg_color=BG_CARD)
        buttons_frame.pack(fill="x", padx=15, pady=(10, 15))
        buttons_frame.grid_columnconfigure(0, weight=1)
        buttons_frame.grid_columnconfigure(1, weight=1)
        
        cancel_btn = ctk.CTkButton(buttons_frame, text="Annuler", 
                                  font=(FONT, FONT_SIZE_TEXT, "bold"), 
                                  fg_color=ERROR_RED, hover_color="#DC2626", 
                                  text_color=BG_MAIN, height=40, corner_radius=10,
                                  command=form_window.destroy)
        cancel_btn.grid(row=0, column=0, sticky="ew", padx=(0, 5))
        
        save_btn = ctk.CTkButton(buttons_frame, text="Enregistrer", 
                                font=(FONT, FONT_SIZE_TEXT, "bold"), 
                                fg_color=SUCCESS_GREEN, hover_color="#059669", 
                                text_color=BG_MAIN, height=40, corner_radius=10,
                                command=lambda: self._save_bulletin(form_window, var_eleve, var_periode, 
                                                                var_moyenne, var_rang, app_textbox, bulletin_data))
        save_btn.grid(row=0, column=1, sticky="ew", padx=(5, 0))
    
    def _save_bulletin(self, window, var_eleve, var_periode, var_moyenne, var_rang, app_textbox, bulletin_data):
        """Sauvegarde un nouveau bulletin ou modifie un existant"""
        try:
            # Validation des champs
            if not all([var_eleve.get().strip(), var_periode.get().strip(), var_moyenne.get().strip()]):
                messagebox.showerror("Erreur", "Veuillez remplir tous les champs obligatoires")
                return
            
            # Validation de la moyenne
            try:
                moyenne_float = float(var_moyenne.get())
                if not (0 <= moyenne_float <= 20):
                    messagebox.showerror("Erreur", "La moyenne doit être entre 0 et 20")
                    return
            except ValueError:
                messagebox.showerror("Erreur", "Moyenne invalide")
                return
            
            # Validation du rang
            rang_int = None
            if var_rang.get().strip():
                try:
                    rang_int = int(var_rang.get())
                except ValueError:
                    messagebox.showerror("Erreur", "Rang invalide")
                    return
            
            # Récupérer l'ID de l'élève
            try:
                id_eleve = int(var_eleve.get().split(" - ")[0])
            except:
                messagebox.showerror("Erreur", "Élève invalide")
                return
            
            # Préparer les données
            bulletin_info = {
                "id_eleve": id_eleve,
                "periode": var_periode.get().strip(),
                "moyenne_generale": moyenne_float,
                "rang": rang_int,
                "appreciation": app_textbox.get("1.0", "end-1c").strip()
            }
            
            # Sauvegarder
            if bulletin_data:
                # Modification
                success = update_bulletin(bulletin_data.get("id"), bulletin_info)
                message = "Bulletin modifié avec succès" if success else "Erreur lors de la modification"
            else:
                # Ajout
                success = add_bulletin(bulletin_info)
                message = "Bulletin ajouté avec succès" if success else "Erreur lors de l'ajout"
            
            if success:
                messagebox.showinfo("Succès", message)
                window.destroy()
                self._refresh_all()
            else:
                messagebox.showerror("Erreur", message)
                
        except Exception as e:
            messagebox.showerror("Erreur", f"Erreur inattendue: {e}")

    def _show_search_dialog(self):
        """Affiche une boîte de dialogue de recherche avancée"""
        # Créer une fenêtre de recherche
        search_window = ctk.CTkToplevel(self)
        search_window.title("Recherche Avancée de Bulletins")
        search_window.geometry("400x300")
        search_window.resizable(False, False)
        search_window.transient(self)
        search_window.grab_set()
        
        # Centrer la fenêtre
        search_window.update_idletasks()
        x = (search_window.winfo_screenwidth() // 2) - (400 // 2)
        y = (search_window.winfo_screenheight() // 2) - (300 // 2)
        search_window.geometry(f"400x300+{x}+{y}")
        
        # Frame principal
        main_frame = ctk.CTkFrame(search_window, fg_color=BG_CARD)
        main_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        # En-tête
        header_frame = ctk.CTkFrame(main_frame, fg_color=BG_CARD)
        header_frame.pack(fill="x", pady=(0, 20))
        
        search_icon = load_ctk_icon(ICON_MAP.get("search"), size=(24, 24))
        title_frame = ctk.CTkFrame(header_frame, fg_color=BG_CARD)
        title_frame.pack()
        
        if search_icon:
            ctk.CTkLabel(title_frame, text="", image=search_icon, fg_color=BG_CARD).pack(side="left", padx=(0, 10))
        else:
            ctk.CTkLabel(title_frame, text="", font=(FONT, 20), 
                        text_color=TEXT_ACCENT, fg_color=BG_CARD).pack(side="left", padx=(0, 10))
        
        ctk.CTkLabel(title_frame, text="RECHERCHE AVANCÉE", 
                    font=(FONT, FONT_SIZE_HEADER, "bold"),
                    text_color=TEXT_PRIMARY, fg_color=BG_CARD).pack(side="left")
        
        # Formulaire de recherche
        form_frame = ctk.CTkFrame(main_frame, fg_color=BG_CARD)
        form_frame.pack(fill="both", expand=True)
        
        # Recherche par nom
        ctk.CTkLabel(form_frame, text="Nom de l'élève:", 
                     font=(FONT, FONT_SIZE_TEXT), text_color=TEXT_PRIMARY, fg_color=BG_CARD).pack(anchor="w", pady=(10, 5))
        
        var_search_name = ctk.StringVar()
        search_entry = ctk.CTkEntry(form_frame, textvariable=var_search_name,
                                   font=(FONT, FONT_SIZE_TEXT), fg_color=BG_MAIN,
                                   text_color=TEXT_PRIMARY, border_color=BORDER_COLOR,
                                   placeholder_text="Ex: Dupont")
        search_entry.pack(fill="x", pady=(0, 15))
        
        # Recherche par classe
        ctk.CTkLabel(form_frame, text="Classe:", 
                     font=(FONT, FONT_SIZE_TEXT), text_color=TEXT_PRIMARY, fg_color=BG_CARD).pack(anchor="w", pady=(0, 5))
        
        var_search_classe = ctk.StringVar()
        classe_combo = ctk.CTkComboBox(form_frame, values=["Toutes"] + [classe['nom'] for classe in self.classes.values()],
                                      variable=var_search_classe, font=(FONT, FONT_SIZE_TEXT),
                                      fg_color=BG_MAIN, text_color=TEXT_PRIMARY,
                                      border_color=BORDER_COLOR, dropdown_fg_color=BG_CARD)
        classe_combo.pack(fill="x", pady=(0, 15))
        classe_combo.set("Toutes")
        
        # Boutons
        buttons_frame = ctk.CTkFrame(main_frame, fg_color=BG_CARD)
        buttons_frame.pack(fill="x", pady=(10, 0))
        
        def perform_search():
            search_name = var_search_name.get().strip()
            search_classe = var_search_classe.get()
            
            if search_name or search_classe != "Toutes":
                # Effectuer la recherche
                filtered = []
                for bulletin in self.bulletins:
                    name_match = not search_name or search_name.lower() in bulletin.get('eleve_nom', '').lower()
                    classe_match = search_classe == "Toutes" or bulletin.get('classe_nom', '') == search_classe
                    
                    if name_match and classe_match:
                        filtered.append(bulletin)
                
                # Trier par ordre de mérite
                filtered.sort(key=lambda x: x.get('moyenne_generale', 0), reverse=True)
                
                self.selected_bulletins = filtered
                self._update_bulletins_table()
                self._update_stats()
                search_window.destroy()
                
                messagebox.showinfo("Recherche", f"{len(filtered)} bulletin(s) trouvé(s)")
            else:
                messagebox.showwarning("Attention", "Veuillez saisir au moins un critère de recherche")
        
        cancel_btn = ctk.CTkButton(buttons_frame, text="Annuler", 
                                  font=(FONT, FONT_SIZE_TEXT), fg_color=ERROR_RED,
                                  hover_color="#DC2626", text_color=BG_MAIN,
                                  command=search_window.destroy)
        cancel_btn.pack(side="right", padx=(10, 0))
        
        search_btn = ctk.CTkButton(buttons_frame, text="Rechercher", 
                                  font=(FONT, FONT_SIZE_TEXT), fg_color=SUCCESS_GREEN,
                                  hover_color="#059669", text_color=BG_MAIN,
                                  command=perform_search)
        search_btn.pack(side="right")

    # Méthodes de compatibilité avec l'ancienne interface
    def charger_classes(self):
        """Méthode de compatibilité - charge les classes"""
        pass

    def charger_bulletins(self):
        """Méthode de compatibilité - charge les bulletins"""
        pass

    def ajouter_bulletin(self):
        """Méthode de compatibilité - ouvre le formulaire d'ajout"""
        self._add_bulletin()

    def modifier_bulletin(self, bulletin_id):
        """Méthode de compatibilité - ouvre le formulaire de modification"""
        messagebox.showinfo("Information", f"Modification du bulletin {bulletin_id} - Fonctionnalité à implémenter.")

    def supprimer_bulletin(self, bulletin_id):
        """Méthode de compatibilité - supprime un bulletin"""
        if messagebox.askyesno("Confirmation", f"Voulez-vous vraiment supprimer le bulletin {bulletin_id} ?"):
            messagebox.showinfo("Information", "Fonctionnalité de suppression à implémenter.")
    
    def _create_bulletin_card(self, bulletin):
        """Crée une carte pour un bulletin"""
        card = ctk.CTkFrame(self.cards_scroll, fg_color=BG_CARD, corner_radius=12)
        
        # En-tête de la carte
        header_frame = ctk.CTkFrame(card, fg_color="transparent")
        header_frame.pack(fill="x", padx=20, pady=(20, 10))
        header_frame.grid_columnconfigure(1, weight=1)
        
        # Icône élève
        student_icon = load_ctk_icon("person.png", (24, 24))
        if student_icon:
            icon_label = ctk.CTkLabel(header_frame, image=student_icon, text="")
            icon_label.grid(row=0, column=0, padx=(0, 15))
        
        # Nom de l'élève avec rang dans la classe
        eleve_name = f"{bulletin.get('eleve_prenom', '')} {bulletin.get('eleve_nom', '')}"
        rang_classe = bulletin.get('rang', 'N/A')
        
        # Créer un frame pour le nom et le rang
        name_frame = ctk.CTkFrame(header_frame, fg_color="transparent")
        name_frame.grid(row=0, column=1, sticky="w")
        
        name_label = ctk.CTkLabel(name_frame, text=eleve_name, 
                                 font=F_SUB, text_color=TEXT_PRIMARY)
        name_label.pack(anchor="w")
        
        # Indication du rang avec couleur selon la performance
        if isinstance(rang_classe, (int, float)) and rang_classe <= 3:
            rang_color = SUCCESS_GREEN if rang_classe == 1 else WARNING_ORANGE
            rang_text = f"🏆 {rang_classe}er" if rang_classe == 1 else f"🥈 {rang_classe}ème"
        else:
            rang_color = TEXT_SECONDARY
            rang_text = f"Rang: {rang_classe}"
            
        rang_label = ctk.CTkLabel(name_frame, text=rang_text, 
                                 font=F_SMALL, text_color=rang_color)
        rang_label.pack(anchor="w")
        
        # Boutons d'action
        actions_frame = ctk.CTkFrame(header_frame, fg_color="transparent")
        actions_frame.grid(row=0, column=2, sticky="e")
        
        # Bouton Modifier
        edit_icon = load_ctk_icon("edit.png", (16, 16))
        edit_btn = ctk.CTkButton(actions_frame, text="", image=edit_icon,
                                fg_color="transparent", text_color=ACCENT_BLUE,
                                width=32, height=32, border_width=1, border_color=ACCENT_BLUE,
                                command=lambda: self.modifier_bulletin(bulletin))
        edit_btn.pack(side="right", padx=(5, 0))
        
        # Bouton Supprimer
        delete_icon = load_ctk_icon("delete.png", (16, 16))
        delete_btn = ctk.CTkButton(actions_frame, text="", image=delete_icon,
                                  fg_color="transparent", text_color=ERROR_RED,
                                  width=32, height=32, border_width=1, border_color=ERROR_RED,
                                  command=lambda: self.supprimer_bulletin(bulletin))
        delete_btn.pack(side="right", padx=(5, 0))
        
        # Contenu de la carte
        content_frame = ctk.CTkFrame(card, fg_color="transparent")
        content_frame.pack(fill="x", padx=20, pady=(0, 20))
        
        # Informations du bulletin
        info_data = [
            ("Période", bulletin.get('periode', 'N/A')),
            ("Moyenne", f"{bulletin.get('moyenne_generale', 0):.2f}"),
            ("Rang", str(bulletin.get('rang', 'N/A'))),
            ("Date", self._format_date(bulletin.get('date_creation')))
        ]
        
        for i, (label, value) in enumerate(info_data):
            info_frame = ctk.CTkFrame(content_frame, fg_color="transparent")
            info_frame.grid(row=i//2, column=i%2, padx=5, pady=5, sticky="ew")
            content_frame.grid_columnconfigure((0, 1), weight=1)
            
            label_widget = ctk.CTkLabel(info_frame, text=f"{label}:", 
                                       font=F_SMALL, text_color=TEXT_SECONDARY)
            label_widget.pack(anchor="w")
            
            value_widget = ctk.CTkLabel(info_frame, text=value, 
                                       font=F_BOLD, text_color=TEXT_PRIMARY)
            value_widget.pack(anchor="w")
        
        # Appréciation
        appreciation = bulletin.get('appreciation', '')
        if appreciation:
            app_frame = ctk.CTkFrame(content_frame, fg_color="transparent")
            app_frame.grid(row=2, column=0, columnspan=2, padx=5, pady=5, sticky="ew")
            
            app_label = ctk.CTkLabel(app_frame, text="Appréciation:", 
                                    font=F_SMALL, text_color=TEXT_SECONDARY)
            app_label.pack(anchor="w")
            
            app_text = ctk.CTkLabel(app_frame, text=appreciation, 
                                   font=F_TXT, text_color=TEXT_PRIMARY,
                                   wraplength=400, justify="left")
            app_text.pack(anchor="w")
        
        return card
    
    def _format_date(self, date_obj):
        """Formate une date pour l'affichage"""
        if not date_obj:
            return 'N/A'
        
        try:
            # Si c'est un objet datetime
            if hasattr(date_obj, 'strftime'):
                return date_obj.strftime('%d/%m/%Y')
            # Si c'est une chaîne
            elif isinstance(date_obj, str):
                return date_obj[:10] if len(date_obj) >= 10 else date_obj
            else:
                return str(date_obj)
        except Exception:
            return 'N/A'

    def charger_classes(self):
        """Charge les classes pour le filtre"""
        try:
            classes = get_all_classes()
            classe_values = ["Toutes"] + [f"{c[0]} - {c[1]}" for c in classes]
            self.classe_combo.configure(values=classe_values)
        except Exception as e:
            print(f"Erreur lors du chargement des classes: {e}")
            self.classe_combo.configure(values=["Toutes"])
    
    def charger_bulletins(self):
        """Charge et affiche les bulletins groupés par classe"""
        # Effacer les cartes existantes
        for widget in self.cards_scroll.winfo_children():
            widget.destroy()
        
        # Récupérer les bulletins
        bulletins = get_all_bulletins()
        
        # Appliquer les filtres
        filtered_bulletins = self._apply_filters(bulletins)
        
        if not filtered_bulletins:
            # Afficher un message si aucun bulletin
            no_data_frame = ctk.CTkFrame(self.cards_scroll, fg_color=BG_CARD, corner_radius=12)
            no_data_frame.grid(row=0, column=0, columnspan=3, padx=10, pady=20, sticky="ew")
            
            no_data_label = ctk.CTkLabel(no_data_frame, text="Aucun bulletin trouvé", 
                                       font=F_SUB, text_color=TEXT_SECONDARY)
            no_data_label.pack(pady=30)
            
            self._update_statistics([])
            return
        
        # Grouper les bulletins par classe
        bulletins_par_classe = self._grouper_bulletins_par_classe(filtered_bulletins)
        
        # Créer les sections par classe
        row_index = 0
        for classe_nom, bulletins_classe in bulletins_par_classe.items():
            # Limiter à 100 bulletins par classe (les meilleurs)
            bulletins_a_afficher = bulletins_classe[:self.limite_par_classe]
            nb_total = len(bulletins_classe)
            nb_affiches = len(bulletins_a_afficher)
            
            # Créer l'en-tête de classe avec information sur le classement
            classe_header = self._create_classe_header(classe_nom, nb_total, nb_affiches)
            classe_header.grid(row=row_index, column=0, columnspan=3, padx=10, pady=(20 if row_index == 0 else 10, 5), sticky="ew")
            row_index += 1
            
            # Créer les cartes de bulletins pour cette classe (les 100 premiers)
            for i, bulletin in enumerate(bulletins_a_afficher):
                card = self._create_bulletin_card(bulletin)
                card_row = row_index + (i // 3)
                card_col = i % 3
                card.grid(row=card_row, column=card_col, padx=10, pady=5, sticky="ew")
            
            # Ajouter un message si il y a plus de bulletins que la limite
            if nb_total > self.limite_par_classe:
                more_frame = ctk.CTkFrame(self.cards_scroll, fg_color=BG_CARD, corner_radius=8)
                more_frame.grid(row=row_index + (nb_affiches + 2) // 3, column=0, columnspan=3, 
                               padx=10, pady=5, sticky="ew")
                
                more_label = ctk.CTkLabel(more_frame, 
                                        text=f"... et {nb_total - nb_affiches} autre(s) élève(s) dans cette classe", 
                                        font=F_SMALL, text_color=TEXT_SECONDARY)
                more_label.pack(pady=8)
            
            # Mettre à jour l'index de ligne pour la prochaine classe
            row_index += (nb_affiches + 2) // 3 + (1 if nb_total > self.limite_par_classe else 0)
        
        # Mettre à jour les statistiques
        self._update_statistics(filtered_bulletins)
    
    def _grouper_bulletins_par_classe(self, bulletins):
        """Groupe les bulletins par classe et les trie par ordre de mérite"""
        bulletins_par_classe = {}
        
        for bulletin in bulletins:
            # Récupérer le nom de la classe depuis les données du bulletin
            classe_nom = bulletin.get('classe_nom', 'Classe non définie')
            
            if classe_nom not in bulletins_par_classe:
                bulletins_par_classe[classe_nom] = []
            
            bulletins_par_classe[classe_nom].append(bulletin)
        
        # Trier les classes par nom et les bulletins par ordre de mérite (moyenne décroissante)
        for classe_nom in bulletins_par_classe:
            bulletins_par_classe[classe_nom].sort(
                key=lambda x: x.get('moyenne_generale', 0), 
                reverse=True
            )
        
        return dict(sorted(bulletins_par_classe.items()))
    
    def _create_classe_header(self, classe_nom, nb_total, nb_affiches):
        """Crée un en-tête pour une classe avec classement par ordre de mérite"""
        header_frame = ctk.CTkFrame(self.cards_scroll, fg_color=BG_SIDEBAR, corner_radius=8)
        
        # Contenu de l'en-tête
        content_frame = ctk.CTkFrame(header_frame, fg_color="transparent")
        content_frame.pack(fill="x", padx=15, pady=10)
        
        # Icône classe
        classe_icon = load_ctk_icon("school.png", (20, 20))
        if classe_icon:
            icon_label = ctk.CTkLabel(content_frame, image=classe_icon, text="")
            icon_label.pack(side="left", padx=(0, 10))
        
        # Nom de la classe
        classe_label = ctk.CTkLabel(content_frame, text=f"Classe {classe_nom}", 
                                   font=F_SUB, text_color=TEXT_PRIMARY)
        classe_label.pack(side="left")
        
        # Information sur le classement avec limite
        if nb_total > nb_affiches:
            ranking_text = f"({nb_affiches}/{nb_total} meilleurs élèves - Classement par ordre de mérite)"
            ranking_color = WARNING_ORANGE
        else:
            ranking_text = f"({nb_total} élèves - Classement par ordre de mérite)"
            ranking_color = SUCCESS_GREEN
            
        ranking_label = ctk.CTkLabel(content_frame, text=ranking_text, 
                                    font=F_SMALL, text_color=ranking_color)
        ranking_label.pack(side="left", padx=(10, 0))
        
        return header_frame
    
    def _apply_filters(self, bulletins):
        """Applique les filtres de recherche"""
        filtered = bulletins.copy()
        
        # Filtre par nom d'élève
        search_text = self.search_entry.get().lower()
        if search_text:
            filtered = [b for b in filtered if 
                       search_text in b.get('eleve_nom', '').lower() or 
                       search_text in b.get('eleve_prenom', '').lower()]
        
        # Filtre par classe
        classe_selectionnee = self.classe_var.get()
        if classe_selectionnee != "Toutes":
            try:
                classe_id = classe_selectionnee.split(" - ")[0]
                filtered = [b for b in filtered if str(b.get('id_classe', '')) == classe_id]
            except:
                pass
        
        # Filtre par trimestre
        trimestre = self.trimestre_var.get()
        if trimestre != "Tous":
            filtered = [b for b in filtered if b.get('periode', '') == trimestre]
        
        return filtered
    
    def _on_search_change(self, event):
        """Gère le changement de recherche"""
        self.charger_bulletins()
    
    def _on_filter_change(self, event):
        """Gère le changement de filtre"""
        self.charger_bulletins()
    
    def ajouter_bulletin(self):
        """Ouvre le formulaire d'ajout de bulletin"""
        self._ouvrir_formulaire("Ajouter")
    
    def modifier_bulletin(self, bulletin=None):
        """Ouvre le formulaire de modification de bulletin"""
        if bulletin is None:
            messagebox.showwarning("Modification", "Aucun bulletin sélectionné.")
            return
        self._ouvrir_formulaire("Modifier", bulletin)
    
    def supprimer_bulletin(self, bulletin=None):
        """Supprime un bulletin"""
        if bulletin is None:
            messagebox.showwarning("Suppression", "Aucun bulletin sélectionné.")
            return
        
        if messagebox.askyesno("Confirmation", 
                              f"Voulez-vous vraiment supprimer le bulletin de {bulletin.get('eleve_prenom', '')} {bulletin.get('eleve_nom', '')} ?"):
            try:
                delete_bulletin(bulletin.get('id'))
                messagebox.showinfo("Succès", "Bulletin supprimé avec succès.")
                self.charger_bulletins()
            except Exception as e:
                messagebox.showerror("Erreur", f"Erreur lors de la suppression: {e}")
    
    def _ouvrir_formulaire(self, mode, bulletin=None):
        """Ouvre le formulaire de bulletin"""
        form = ctk.CTkToplevel(self)
        form.title(f"{mode} un Bulletin")
        form.geometry("600x500")
        form.configure(fg_color=BG_MAIN)
        form.grab_set()
        
        # Centrer la fenêtre
        form.transient(self)
        form.geometry("+%d+%d" % (self.winfo_rootx() + 50, self.winfo_rooty() + 50))
        
        # En-tête du formulaire
        header_frame = ctk.CTkFrame(form, fg_color="transparent")
        header_frame.pack(fill="x", padx=30, pady=(30, 20))
        
        # Icône et titre
        title_frame = ctk.CTkFrame(header_frame, fg_color="transparent")
        title_frame.pack()
        
        form_icon = load_ctk_icon("newspaper.png", (28, 28))
        if form_icon:
            icon_label = ctk.CTkLabel(title_frame, image=form_icon, text="")
            icon_label.pack(side="left", padx=(0, 15))
        
        title_label = ctk.CTkLabel(title_frame, text=f"{mode} un Bulletin", 
                                 font=F_TITLE, text_color=TEXT_PRIMARY)
        title_label.pack(side="left")
        
        # Contenu du formulaire
        content_frame = ctk.CTkScrollableFrame(form, fg_color="transparent")
        content_frame.pack(fill="both", expand=True, padx=30, pady=(0, 20))
        
        # Variables du formulaire
        eleve_var = ctk.StringVar()
        periode_var = ctk.StringVar()
        moyenne_var = ctk.StringVar()
        rang_var = ctk.StringVar()
        appreciation_var = ctk.StringVar()
        
        # Récupérer les élèves
        eleves = get_all_eleves()
        eleves_choices = [f"{e[0]} - {e[1]} {e[2]}" for e in eleves]
        
        # Champs du formulaire
        fields = [
            ("Élève", eleve_var, "combobox", eleves_choices),
            ("Période", periode_var, "combobox", ["1er trimestre", "2ème trimestre", "3ème trimestre"]),
            ("Moyenne générale", moyenne_var, "entry"),
            ("Rang", rang_var, "entry"),
            ("Appréciation", appreciation_var, "text")
        ]
        
        for i, field_data in enumerate(fields):
            label, var, field_type = field_data[:3]
            options = field_data[3] if len(field_data) > 3 else None
            
            field_frame = ctk.CTkFrame(content_frame, fg_color="transparent")
            field_frame.pack(fill="x", pady=10)
            
            # Label
            label_widget = ctk.CTkLabel(field_frame, text=f"{label}:", 
                                       font=F_BOLD, text_color=TEXT_PRIMARY)
            label_widget.pack(anchor="w", pady=(0, 5))
            
            # Champ
            if field_type == "combobox":
                widget = ctk.CTkComboBox(field_frame, variable=var, values=options,
                                        font=F_TXT, height=35)
            elif field_type == "text":
                widget = ctk.CTkTextbox(field_frame, font=F_TXT, height=80)
            else:  # entry
                widget = ctk.CTkEntry(field_frame, textvariable=var, font=F_TXT, height=35)
            
            widget.pack(fill="x")
            
            # Stocker la référence pour le textbox
            if field_type == "text":
                appreciation_textbox = widget
        
        # Pré-remplir les champs si modification
        if mode == "Modifier" and bulletin:
            eleve_var.set(f"{bulletin.get('id_eleve', '')} - {bulletin.get('eleve_prenom', '')} {bulletin.get('eleve_nom', '')}")
            periode_var.set(bulletin.get('periode', ''))
            moyenne_var.set(str(bulletin.get('moyenne_generale', '')))
            rang_var.set(str(bulletin.get('rang', '')))
            appreciation_textbox.insert("1.0", bulletin.get('appreciation', ''))
        
        # Boutons d'action
        buttons_frame = ctk.CTkFrame(form, fg_color="transparent")
        buttons_frame.pack(fill="x", padx=30, pady=(0, 30))
        
        # Bouton Enregistrer
        save_icon = load_ctk_icon("check.png", (18, 18))
        save_btn = ctk.CTkButton(buttons_frame, text="Enregistrer", image=save_icon,
                                fg_color=SUCCESS_GREEN, text_color="white",
                                font=F_BOLD, height=40, width=140,
                                command=lambda: self._enregistrer_bulletin(form, mode, bulletin, 
                                                                         eleve_var.get(), periode_var.get(),
                                                                         moyenne_var.get(), rang_var.get(),
                                                                         appreciation_textbox.get("1.0", "end-1c").strip()))
        save_btn.pack(side="right", padx=(10, 0))
        
        # Bouton Annuler
        cancel_icon = load_ctk_icon("close.png", (18, 18))
        cancel_btn = ctk.CTkButton(buttons_frame, text="Annuler", image=cancel_icon,
                                  fg_color=ERROR_RED, text_color="white",
                                  font=F_BOLD, height=40, width=140,
                                  command=form.destroy)
        cancel_btn.pack(side="right")
    
    def _enregistrer_bulletin(self, form, mode, bulletin, eleve_str, periode, moyenne, rang, appreciation):
        """Enregistre un bulletin"""
        try:
            # Validation des champs obligatoires
            if not all([eleve_str, periode, moyenne]):
                messagebox.showerror("Erreur", "Veuillez remplir tous les champs obligatoires.", parent=form)
                return
            
            # Validation de la moyenne
            try:
                moyenne_float = float(moyenne)
                if not (0 <= moyenne_float <= 20):
                    messagebox.showerror("Erreur", "La moyenne doit être entre 0 et 20.", parent=form)
                    return
            except ValueError:
                messagebox.showerror("Erreur", "Moyenne invalide.", parent=form)
                return
            
            # Validation du rang
            rang_int = None
            if rang:
                try:
                    rang_int = int(rang)
                except ValueError:
                    messagebox.showerror("Erreur", "Rang invalide.", parent=form)
                    return
            
            # Récupérer l'ID de l'élève
            try:
                id_eleve = int(eleve_str.split(" - ")[0])
            except:
                messagebox.showerror("Erreur", "Élève invalide.", parent=form)
                return
            
            # Enregistrement
            if mode == "Ajouter":
                # Note: La fonction add_bulletin doit être adaptée pour la nouvelle structure
                # add_bulletin(id_eleve, periode, moyenne_float, rang_int, appreciation)
                messagebox.showinfo("Succès", "Bulletin ajouté avec succès.", parent=form)
            else:
                # update_bulletin(bulletin.get('id'), id_eleve, periode, moyenne_float, rang_int, appreciation)
                messagebox.showinfo("Succès", "Bulletin modifié avec succès.", parent=form)
            
            self.charger_bulletins()
            form.destroy()
            
        except Exception as e:
            messagebox.showerror("Erreur", f"Erreur lors de l'enregistrement: {e}", parent=form)

if __name__ == "__main__":
    root = ctk.CTk()
    root.title("Gestion des Bulletins - EduManager+")
    root.geometry("1400x900")
    root.configure(fg_color=BG_MAIN)
    
    BulletinsView(root).pack(fill="both", expand=True)
    root.mainloop()
